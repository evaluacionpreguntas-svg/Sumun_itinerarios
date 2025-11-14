import streamlit as st
import pandas as pd
import vertexai
from vertexai.preview.generative_models import GenerativeModel
from vertexai.preview.vision_models import ImageGenerationModel
from vertexai.preview.generative_models import Part
import PyPDF2
import docx
import re
import io
import os
import json
import zipfile
# --- NUEVA IMPORTACIÓN PARA GCS ---
from google.cloud import storage
import requests
from streamlit_lottie import st_lottie
from graficos_plugins import generar_grafico_desde_texto
from docx.shared import Inches
import random
from dotenv import load_dotenv
import numpy as np
from langchain_text_splitters import RecursiveCharacterTextSplitter
from vertexai.language_models import TextEmbeddingModel # Para los embeddings
import fitz  # PyMuPDF


def parse_json_llm(s: str):
    """
    Busca y decodifica un objeto JSON dentro de un string, 
    limpiando errores comunes de formato de los LLMs.
    Devuelve un diccionario si tiene éxito, o None si no puede.
    """
    if not s:
        return None
    
    # Quitar cercos de código tipo ```json ... ``` o ``` ... ```
    s = re.sub(r"^```(?:json)?\s*|\s*```$", "", s.strip(), flags=re.DOTALL)

    # Buscar el objeto JSON principal (del primer '{' al último '}')
    start = s.find('{')
    end = s.rfind('}')
    if start == -1 or end == -1:
        return None
        
    json_str = s[start:end+1]

    try:
        # Intentar decodificar
        return json.loads(json_str)
    except json.JSONDecodeError:
        # Si falla, podría estar "doble codificado" (un string JSON dentro de otro)
        try:
            decoded_str = json.loads(f'"{json_str}"')
            return json.loads(decoded_str)
        except Exception:
            return None

def load_bloom_taxonomy(file_path="bloom_taxonomy.json"):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        st.error(f"Error: No se encontró el archivo de la taxonomía en '{file_path}'.")
        return {}
    except json.JSONDecodeError:
        st.error(f"Error: El archivo '{file_path}' no es un JSON válido.")
        return {}

# Carga la taxonomía al inicio de tu script
bloom_taxonomy_detallada = load_bloom_taxonomy()

# --- REEMPLAZA ESTA FUNCIÓN ---
def crear_indice_vectorial(paginas_texto):
    """
    Convierte una lista de TEXTOS DE PÁGINA en chunks y vectores.
    Procesa página por página para ahorrar RAM.
    """
    try:
        model = TextEmbeddingModel.from_pretrained("text-embedding-004")
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100
        )
        
        index = []
        
        # --- AQUÍ ESTÁ EL CAMBIO ---
        # Reducimos el tamaño del lote de 250 a 100.
        # Esto asegura que no superemos el límite de 20k tokens por llamada.
        api_batch_size = 100 
        # --- FIN DEL CAMBIO ---

        # Iteramos sobre CADA PÁGINA individualmente
        for texto_pagina in paginas_texto:
            # 1. Dividimos solo el texto de ESTA página
            chunks_pagina = text_splitter.split_text(texto_pagina)
            
            if not chunks_pagina:
                continue
            
            # 2. Vectorizamos los chunks de ESTA página (en lotes si es necesario)
            for i in range(0, len(chunks_pagina), api_batch_size):
                batch_chunks = chunks_pagina[i:i + api_batch_size]
                embeddings = model.get_embeddings(batch_chunks)
                
                for chunk, embedding in zip(batch_chunks, embeddings):
                    index.append((chunk, np.array(embedding.values)))
            
            # 3. Al final de este bucle, 'texto_pagina' y 'chunks_pagina' se liberan
            # de la memoria antes de procesar la siguiente página.
    
        return index
    
    except Exception as e:
        # ¡Este es el error que estás viendo ahora!
        st.error(f"Error al crear vectores (Embeddings): {e}")
        return []

def buscar_en_indice(query_text, k=3):
        """
        Busca en el índice de sesión los k chunks más relevantes para un texto.
        Devuelve una lista de los textos (chunks) encontrados.
        """
        if 'pdf_index' not in st.session_state or not st.session_state['pdf_index']:
            return [] # No hay índice cargado

        index = st.session_state['pdf_index']
        
        try:
            # 1. Vectorizar la consulta (la microhabilidad)
            model = TextEmbeddingModel.from_pretrained("text-embedding-004")
            query_vector = np.array(model.get_embeddings([query_text])[0].values)
            
            # 2. Calcular similitud (Coseno)
            # Preparamos los vectores del índice
            chunk_vectors = np.array([item[1] for item in index])
            
            # Normalizamos vectores
            query_norm = np.linalg.norm(query_vector)
            chunk_norms = np.linalg.norm(chunk_vectors, axis=1)
            
            # Evitar división por cero si hay vectores nulos
            if query_norm == 0 or np.any(chunk_norms == 0):
                return []
                
            # Calculamos la similitud del coseno
            # (A . B) / (||A|| * ||B||)
            similitudes = np.dot(chunk_vectors, query_vector) / (chunk_norms * query_norm)
            
            # 3. Obtener los Top K
            # `np.argsort` da los índices de menor a mayor. Usamos `[-k:]` para los k más altos
            # y `[::-1]` para invertirlos (del más al menos relevante).
            top_k_indices = np.argsort(similitudes)[-k:][::-1]
            
            # 4. Devolver los textos de esos chunks
            relevant_chunks = [index[i][0] for i in top_k_indices]
            return relevant_chunks

        except Exception as e:
            st.warning(f"Error al buscar en el índice del PDF: {e}")
            return []
    

def load_lottieurl(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

def _pick(d, *keys):
    for k in keys:
        v = d.get(k)
        if v not in (None, ""):
            return v
    return None
    
def normaliza_claves_classif(c: dict) -> dict:
    if not c:
        return {}
    return {
        'ID': _pick(c, 'ID', 'Id', 'id'),
        'GRADO': _pick(c, 'GRADO', 'Grado', 'grado'),
        'ÁREA': _pick(c, 'ÁREA', 'Área', 'Area', 'área', 'area'),
        'ASIGNATURA': _pick(c, 'ASIGNATURA', 'Asignatura', 'asignatura'),
        'MACROHABILIDAD': _pick(c, 'MACROHABILIDAD', 'Macrohabilidad', 'Macrohabilidad', 'macrohabilidad', 'macrohabilidad'),
        'PROCESO COGNITIVO': _pick(c, 'PROCESO COGNITIVO', 'Proceso Cognitivo', 'proceso cognitivo'),
        'MICROHABILIDAD': _pick(c, 'MICROHABILIDAD', 'Microhabilidad'),
        'COMPETENCIA MICROHABILIDAD': _pick(c, 'COMPETENCIA MICROHABILIDAD', 'Competencia Microhabilidad'),
        'Numero': _pick(c, 'Numero', 'NUMERO', 'Número', 'NÚMERO', 'numero', 'número'),
    }

def describir_imagen_con_llm(model_name, image_bytes, file_type):
    """
    Usa un modelo multimodal de Vertex AI para generar una descripción detallada de una imagen.
    """
    try:
        # Asegúrate de usar un modelo que soporte visión, como gemini-1.5-pro
        model = GenerativeModel(model_name)
        
        # Prepara la imagen para el modelo
        image_part = Part.from_data(data=image_bytes, mime_type=file_type)

        # Prompt específico para obtener una descripción útil para evaluaciones
        prompt_descripcion = """
        Describe esta imagen con el máximo nivel de detalle posible, como si se la estuvieras describiendo a alguien que no puede verla y necesita construir una pregunta de evaluación sobre ella.
        Enfócate en los siguientes aspectos:
        1.  **Objetos y Entidades:** Lista todos los objetos, personas o animales presentes.
        2.  **Acciones y Relaciones:** Describe qué está ocurriendo y cómo interactúan los elementos entre sí.
        3.  **Texto y Símbolos:** Transcribe cualquier texto visible, números, etiquetas o símbolos importantes.
        4.  **Composición y Contexto:** Describe la escena general, la disposición de los elementos y cualquier inferencia obvia sobre el lugar o la situación.
        Genera una descripción completa en un único párrafo.
        """
        
        # Genera el contenido
        response = model.generate_content([prompt_descripcion, image_part])
        return response.text
    except Exception as e:
        st.error(f"Error al describir la imagen con Vertex AI: {e}")
        return None


def extraer_texto_pdf(pdf_bytes):
    """
    Extrae texto de un PDF en bytes USANDO PyMuPDF (fitz),
    que es mucho más rápido y robusto que PyPDF2.
    """
    try:
        texto_pdf = ""
        # Abrir el PDF desde los bytes en memoria
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            for page in doc:
                # .get_text() es el método de PyMuPDF
                texto_pdf += page.get_text() + "\n\n" 
        return texto_pdf
    except Exception as e:
        # Aún mantenemos el st.error por si el PDF está dañado
        st.error(f"Error al leer el PDF con PyMuPDF: {e}")
        return None



# --- FUNCIÓN PRINCIPAL QUE ENVUELVE TODA LA APP ---
def main():
    # --- CONFIGURACIÓN DE LA PÁGINA DE STREAMLIT ---
    st.set_page_config(
        page_title="Generador y Auditor de Ítems con IA (Vertex AI)",
        page_icon="🧠",
        layout="wide"
    )

    # --- INICIALIZACIÓN DE VERTEX AI ---
    try:
        GCP_PROJECT_ID = os.environ.get("GCP_PROJECT_ID")
        GCP_LOCATION = os.environ.get("GCP_LOCATION")
        vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
        st.sidebar.success("✅ Conectado a Vertex AI.")
    except Exception as e:
        st.sidebar.error(f"Error al inicializar Vertex AI: {e}")
        st.error("No se pudo conectar con Vertex AI. Verifica la configuración del proyecto y la autenticación.")
        st.stop()

    # --- Título Principal de la App ---
    st.title("📚 Generador y Auditor de ítems final de itinerario SUMUN 🧠")
    st.markdown("Esta aplicación genera ítems de selección múltiple y audita su calidad utilizando modelos de **Google Cloud Vertex AI**.")


    definiciones_tipologias = {
        "Crónica": "Un relato narrativo y detallado de hechos, usualmente presentados en orden cronológico. Debe incluir la perspectiva y estilo subjetivo del autor, mezclando información con impresiones personales.",
        "Noticia": "Un texto objetivo y conciso que informa sobre un evento reciente y de interés público. Debe responder a las preguntas qué, quién, cómo, cuándo, dónde y por qué. El lenguaje debe ser formal y directo.",
        "Entrevista": "Un texto que presenta un diálogo entre un entrevistador y un entrevistado. Debe estar en formato de pregunta y respuesta (Ej: 'Entrevistador: ...', 'Entrevistado: ...') y revelar información u opiniones del entrevistado.",
        "Ensayo": "Un texto en prosa que analiza, interpreta o evalúa un tema desde una perspectiva personal y argumentativa. Debe presentar una tesis clara y desarrollarla con argumentos y reflexiones.",
        "Cuento Corto": "Un relato de ficción breve, con pocos personajes y una trama concisa que se desarrolla hacia un clímax y un final. Debe tener elementos narrativos claros como inicio, nudo y desenlace.",
        "Manual": "Un texto instructivo y funcional que explica paso a paso cómo realizar una tarea o usar un producto. El lenguaje debe ser claro, preciso y directo, a menudo usando listas numeradas o viñetas."
    }
    
    # -------------------------------------------------------------------
    # --- SECCIÓN DE DEFINICIÓN DE TODAS LAS FUNCIONES DE AYUDA ---
    # -------------------------------------------------------------------
    
    ### INICIO DE NUEVAS FUNCIONALIDADES DE AUTOGUARDADO ###
    def generar_nombre_archivo_progreso(grado, asignatura, macrohabilidad):
        """Crea un nombre de archivo seguro y único basado en las selecciones."""
        grado_str = str(grado).replace(" ", "_")
        asignatura_str = str(asignatura).replace(" ", "_")
        macrohabilidad_str = str(macrohabilidad).replace(" ", "_")
        nombre_base = f"progreso_{grado_str}_{asignatura_str}_{macrohabilidad_str}"
        nombre_seguro = re.sub(r'[^a-zA-Z0-9_.-]', '', nombre_base)
        return f"{nombre_seguro}.json"

    def guardar_progreso_en_gcs(bucket_name, file_name, data):
        """Guarda el estado de la sesión en un archivo JSON en GCS."""
        if not bucket_name: return
        try:
            storage_client = storage.Client()
            bucket = storage_client.bucket(bucket_name)
            blob = bucket.blob(f"progreso/{file_name}") 
            json_data = json.dumps(data, indent=4)
            blob.upload_from_string(json_data, content_type='application/json')
        except Exception as e:
            st.sidebar.warning(f"No se pudo autoguardar el progreso: {e}")

    def cargar_progreso_desde_gcs(bucket_name, file_name):
        """Carga el estado de la sesión desde GCS, si existe."""
        if not bucket_name: return []
        try:
            storage_client = storage.Client()
            bucket = storage_client.bucket(bucket_name)
            blob = bucket.blob(f"progreso/{file_name}")
            if blob.exists():
                json_data = blob.download_as_string()
                data = json.loads(json_data)
                st.sidebar.success(f"Progreso recuperado para esta macrohabilidad.")
                return data
            return []
        except Exception as e:
            st.sidebar.error(f"Error al cargar el progreso: {e}")
            return []

    def borrar_progreso_en_gcs(bucket_name, file_name):
        """Borra el archivo de progreso de GCS al reiniciar."""
        if not bucket_name: return
        try:
            storage_client = storage.Client()
            bucket = storage_client.bucket(bucket_name)
            blob = bucket.blob(f"progreso/{file_name}")
            if blob.exists():
                blob.delete()
                st.sidebar.info("El progreso guardado ha sido eliminado.")
        except Exception as e:
            st.sidebar.warning(f"No se pudo borrar el progreso: {e}")
    ### FIN DE NUEVAS FUNCIONALIDADES ###
    
    @st.cache_data
    def leer_excel_desde_gcs(bucket_name, file_path):
        """
        Lee un archivo Excel directamente desde un bucket de GCS.
        """
        try:
            storage_client = storage.Client()
            bucket = storage_client.bucket(bucket_name)
            blob = bucket.blob(file_path)
            
            file_bytes = blob.download_as_bytes()
            
            df = pd.read_excel(io.BytesIO(file_bytes))
            st.sidebar.success(f"Archivo Excel '{file_path}' cargado desde GCS.")
            return df
        except Exception as e:
            st.sidebar.error(f"Error al leer Excel desde GCS: {e}")
            st.error(f"No se pudo cargar el archivo Excel desde el bucket '{bucket_name}'. Revisa los permisos y la ruta del archivo.")
            return None

    @st.cache_data
    def leer_pdf_desde_gcs(bucket_name, file_path):
        """
        Lee el texto de un archivo PDF directamente desde un bucket de GCS.
        """
        try:
            storage_client = storage.Client()
            bucket = storage_client.bucket(bucket_name)
            blob = bucket.blob(file_path)
            
            file_bytes = blob.download_as_bytes()
            
            texto_pdf = ""
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                texto_pdf += page.extract_text() or ""
            
            st.sidebar.success(f"Archivo PDF '{file_path}' cargado desde GCS.")
            return texto_pdf
        except Exception as e:
            st.sidebar.error(f"Ocurrió un error al leer el PDF desde GCS: {e}")
            st.error(f"No se pudo cargar el archivo PDF desde el bucket '{bucket_name}'. Revisa los permisos y la ruta del archivo.")
            return ""

    # --- CÓDIGO DE REEMPLAZO ---
    
    def generar_contexto_general_con_llm(model_name, grado, area, asignatura, macrohabilidad, tipo_contexto="", idea_usuario=""):
        """
        Genera un texto de contexto general para una macrohabilidad, aplicando una tipología textual específica si se proporciona.
        """
        # Construimos una sección especial para la tipología solo si se especifica una
        seccion_tipologia = ""
        if tipo_contexto in definiciones_tipologias:
            definicion = definiciones_tipologias[tipo_contexto]
            seccion_tipologia = f"""
    --- TIPOLOGÍA TEXTUAL OBLIGATORIA ---
    ¡INSTRUCCIÓN CRÍTICA! El texto generado DEBE corresponder fielmente a la siguiente tipología:
    - TIPO: {tipo_contexto}
    - DEFINICIÓN: {definicion}
    Garantiza que la estructura, el estilo y el lenguaje del texto cumplan con esta definición.
    ------------------------------------
    """
    
        prompt_contexto = f'''Eres un escritor y diseñador instruccional experto. Tu tarea es redactar un texto de CONTEXTO para una evaluación educativa.
    
    --- PARÁMETROS GENERALES ---
    - Grado: {grado}
    - Área: {area}
    - Asignatura: {asignatura}
    - Macrohabilidad o unidad temática: {macrohabilidad}
    {seccion_tipologia}
    --- IDEA GUÍA DEL USUARIO (Opcional) ---
    {idea_usuario if idea_usuario else "No se proporcionó una idea específica, usa tu creatividad para el tema."}
    ------------------------------------
    
    --- INSTRUCCIONES FINALES ---
    - El texto debe ser coherente, verosímil y apropiado para el nivel educativo.
    - Si se especificó una tipología, CÚMPLELA rigurosamente.
    - El texto final debe tener entre 150 y 300 palabras.
    - Devuelve ÚNICAMENTE el texto del contexto. NO incluyas títulos (a menos que la tipología lo requiera, como en una noticia), preguntas, ni explicaciones.
    '''
        try:
            modelo = GenerativeModel(model_name)
            response = modelo.generate_content(prompt_contexto)
            return response.text
        except Exception as e:
            st.error(f"Error al generar el contexto con Vertex AI: {e}")
            return None
    
    def refinar_contexto_con_llm(model_name, contexto_original, feedback_usuario):
        """
        Refina un texto de contexto existente basado en el feedback del usuario.
        """
        prompt_refinamiento = f"""
    Eres un editor experto. Tu tarea es reescribir y mejorar el siguiente texto de CONTEXTO basado en las observaciones del usuario.
    No cambies la intención original del texto a menos que el feedback te lo pida. El objetivo es ajustar y perfeccionar.
    
    --- TEXTO ORIGINAL ---
    {contexto_original}
    --------------------
    
    --- OBSERVACIONES DEL USUARIO PARA REFINAR ---
    {feedback_usuario}
    -------------------------------------------
    
    --- INSTRUCCIONES ---
    - Devuelve únicamente el texto del contexto refinado.
    - No añadas explicaciones, saludos ni ningún otro texto fuera del contexto en sí.
    """
        try:
            modelo = GenerativeModel(model_name)
            response = modelo.generate_content(prompt_refinamiento)
            return response.text
        except Exception as e:
            st.error(f"Error al refinar el contexto con Vertex AI: {e}")
            return None
            
    def get_descripcion_detallada_bloom(proceso_cognitivo_elegido):
        """
        Busca el proceso cognitivo en la estructura detallada y devuelve una descripción formateada.
        """
        proceso_upper = str(proceso_cognitivo_elegido).upper()
        proceso_data = bloom_taxonomy_detallada.get(proceso_upper)

        if not proceso_data:
            return "Descripción no disponible para este proceso cognitivo."

        descripcion_formateada = f"**Categoría Cognitiva: {proceso_upper}**\n"
        descripcion_formateada += f"- **Definición General**: {proceso_data['definicion']}\n\n"
        descripcion_formateada += "**Subprocesos Cognitivos Asociados:**\n"

        for subproceso, detalles in proceso_data['subprocesos'].items():
            descripcion_formateada += f"- **{subproceso}** (Nombres alternativos: {detalles['nombres_alternativos']}):\n"
            descripcion_formateada += f"  - {detalles['definicion_ejemplo']}\n"

        return descripcion_formateada
        
    def get_preguntas_tipo_formateadas(proceso_cognitivo_elegido):
        """
        Busca un proceso cognitivo y devuelve un string formateado 
        con todas las preguntas tipo de sus subprocesos.
        """
        proceso_upper = str(proceso_cognitivo_elegido).upper()
        proceso_data = bloom_taxonomy_detallada.get(proceso_upper)
    
        if not proceso_data or 'subprocesos' not in proceso_data:
            return "No se encontraron ejemplos de preguntas para este proceso."
    
        texto_formateado = ""
        for subproceso, detalles in proceso_data['subprocesos'].items():
            if 'Preguntas_tipo' in detalles and detalles['Preguntas_tipo']:
                texto_formateado += f"\nPara el subproceso '{subproceso}':\n"
                for pregunta in detalles['Preguntas_tipo']:
                    texto_formateado += f"- {pregunta}\n"
        
        return texto_formateado
            
    
    def generar_texto_con_llm(model_name, prompt, force_json=False):
        """
        Genera texto usando un modelo de Vertex AI, con opción para forzar salida JSON.
        """
        try:
            modelo = GenerativeModel(model_name)
            gen_config = {}
            if force_json:
                gen_config["response_mime_type"] = "application/json"
                
            response = modelo.generate_content(prompt, generation_config=gen_config)
            return response.text
        except Exception as e:
            st.error(f"Error al llamar a Vertex AI ({model_name}): {e}")
            return None

    def auditar_item_con_llm(model_name, item_generado, grado, area, asignatura, macrohabilidad,
                             proceso_cognitivo, microhabilidad,
                             competencia_microhabilidad, contexto_educativo, manual_reglas_texto="", 
                             descripcion_bloom="", grafico_necesario="", descripcion_grafico="", 
                             prompt_auditor_adicional=""):
        """
        Audita un ítem generado para verificar su cumplimiento con criterios específicos,
        solicitando y esperando una respuesta en formato JSON.
        """
        auditoria_prompt = f"""
    Eres un experto en validación de ítems educativos, especializado en pruebas tipo ICFES y las directrices del equipo IMPROVE.
    Tu tarea es AUDITAR RIGUROSAMENTE el siguiente ítem generado por un modelo de lenguaje y devolver tu análisis en formato JSON.
    Debes verificar que el ítem cumpla con TODOS los siguientes criterios, prestando especial atención a la alineación con los parámetros proporcionados y a las reglas de formato y contenido.
    
    --- CRITERIOS DE AUDITORÍA ---
    Evalúa el ítem de manera rigurosa según los siguientes puntos clave. La calidad de tu auditoría depende de la atención a cada detalle.
    
    1.  **Formato del Enunciado:** La pregunta debe ser clara, directa, sin ambigüedades ni errores gramaticales.
    
    2.  **Estilo del Enunciado (Regla de No Jerarquización):** ¡CRITERIO CRÍTICO! Verifica que el enunciado no contenga palabras comparativas o superlativas que obliguen al estudiante a jerarquizar las opciones (ej. "más", "mejor", "principal"). La violación de esta regla es un error crítico y debe resultar en un **RECHAZO** automático en el dictamen final.
    
    3.  **Número de Opciones:** Debe haber exactamente 4 opciones (A, B, C, D).
    
    4.  **Respuesta Correcta Indicada:** La sección 'RESPUESTA CORRECTA:' debe estar presente, claramente indicada y coincidir con una de las opciones.
    
    5.  **Diseño de Justificaciones:** Deben existir justificaciones diferenciadas para cada opción. La de la opción correcta debe explicar el razonamiento (no por descarte), y las de las incorrectas deben seguir el formato: “El estudiante podría escoger… Sin embargo, esto es incorrecto porque…”.
    
    6.  **Estilo y Restricciones:** El ítem debe evitar negaciones mal redactadas, nombres/marcas/lugares reales, datos personales y frases vagas como “ninguna de las anteriores” o “todas las anteriores”.
    
    7.  **Alineación del Contenido:** Evalúa si el ítem se alinea EXCLUSIVAMENTE con todos los parámetros:
        * **Temáticos:** Grado (`{grado}`), Área (`{area}`), Asignatura (`{asignatura}`), Macrohabilidad (`{macrohabilidad}`), Microhabilidad (`{microhabilidad}`).
        * **Cognitivos:** Proceso (`{proceso_cognitivo}`). Crucialmente, verifica la **exclusividad cognitiva**: la tarea debe ser demostrablemente más compleja que el nivel cognitivo anterior y no debe requerir un nivel superior.
    
    8.  **Gráfico (si aplica):** Si se requiere un gráfico (`{grafico_necesario}`), la descripción (`{descripcion_grafico}`) debe ser clara y funcional.
    
    --- MANUAL DE REGLAS ADICIONAL ---
    {manual_reglas_texto}
    -----------------------------------
    
    --- INSTRUCCIONES ADICIONALES PARA LA AUDITORÍA ---
    {prompt_auditor_adicional if prompt_auditor_adicional else "No se proporcionaron instrucciones adicionales para la auditoría."}
    ---------------------------------------------------
    
    --- ÍTEM A AUDITAR ---
    {item_generado}
    --------------------
    
    ¡INSTRUCCIÓN CRÍTICA DE SALIDA!
    Devuelve tu auditoría como un único bloque de código JSON válido, sin ningún otro texto o explicación antes o después. No uses ```json.
    El objeto JSON debe tener la siguiente estructura exacta, incluyendo un objeto por cada uno de los 8 criterios:
    {{
      "criterios": [
        {{
          "criterio": "Formato del Enunciado",
          "estado": "✅",
          "comentario": "El enunciado es claro y directo."
        }},
        {{
          "criterio": "Estilo del Enunciado (Regla de No Jerarquización)",
          "estado": "✅",
          "comentario": "El enunciado no utiliza términos comparativos."
        }},
        {{
          "criterio": "Número de Opciones",
          "estado": "✅",
          "comentario": "Se presentan 4 opciones."
        }},
        {{
          "criterio": "Respuesta Correcta Indicada",
          "estado": "✅",
          "comentario": "La clave de respuesta está presente y es válida."
        }},
        {{
          "criterio": "Diseño de Justificaciones",
          "estado": "✅",
          "comentario": "Las justificaciones siguen el formato requerido."
        }},
        {{
          "criterio": "Estilo y Restricciones",
          "estado": "✅",
          "comentario": "El ítem no contiene nombres propios, marcas u otras restricciones."
        }},
        {{
          "criterio": "Alineación del Contenido",
          "estado": "✅",
          "comentario": "El ítem se alinea correctamente con los parámetros temáticos y cognitivos."
        }},
        {{
          "criterio": "Gráfico (si aplica)",
          "estado": "✅",
          "comentario": "No se requiere gráfico, lo cual es correcto para este ítem."
        }}
      ],
      "dictamen_final": "✅ CUMPLE TOTALMENTE",
      "observaciones_finales": "El ítem cumple con todos los criterios de auditoría y se considera apto."
    }}
    """
        return generar_texto_con_llm(model_name, auditoria_prompt, force_json=True), auditoria_prompt

               
    def generar_pregunta_con_seleccion(gen_model_name, audit_model_name,
                                         fila_datos, criterios_generacion, manual_reglas_texto="",
                                         informacion_adicional_usuario="",
                                         prompt_bloom_adicional="", prompt_construccion_adicional="", prompt_especifico_adicional="",
                                         prompt_auditor_adicional="",
                                         contexto_general_macrohabilidad="", contexto_del_libro="",  feedback_usuario="", item_a_refinar_text="", descripcion_imagen_aprobada=""):
        """
        Genera una pregunta educativa de opción múltiple usando el modelo de generación seleccionado
        y la itera para refinarla si la auditoría lo requiere.
        """
        tipo_pregunta = criterios_generacion.get("tipo_pregunta", "opción múltiple con 4 opciones")
        dificultad = criterios_generacion.get("dificultad", "media")
        contexto_educativo = criterios_generacion.get("contexto_educativo", "general")
        formato_justificacion = criterios_generacion.get("formato_justificacion", """
            • Justificación correcta: debe explicar el razonamiento o proceso cognitivo (NO por descarte).
            • Justificaciones incorrectas: deben redactarse como: “El estudiante podría escoger la opción X porque… Sin embargo, esto es incorrecto porque…”
        """)
        fila_datos = normaliza_claves_classif(fila_datos)
        grado_elegido = fila_datos.get('GRADO', 'no especificado')
        area_elegida = fila_datos.get('ÁREA', 'no especificada')
        asignatura_elegida = fila_datos.get('ASIGNATURA', 'no especificada')
        macrohabilidad_elegida = fila_datos.get('MACROHABILIDAD', 'no especificada')
        proceso_cognitivo_elegido = fila_datos.get('PROCESO COGNITIVO', 'no especificado')
        microhabilidad_elegida = fila_datos.get('MICROHABILIDAD', 'no especificada')
        competencia_microhabilidad_elegida = fila_datos.get('COMPETENCIA MICROHABILIDAD', 'no especificada')
    
        dato_para_pregunta_foco = microhabilidad_elegida
        descripcion_bloom = get_descripcion_detallada_bloom(proceso_cognitivo_elegido)
        ejemplos_preguntas = get_preguntas_tipo_formateadas(proceso_cognitivo_elegido)
        
        current_item_text = ""
        auditoria_status = "❌ RECHAZADO"
        audit_observations = ""
        max_refinement_attempts = 3
        attempt = 0
        grafico_necesario = "NO"
        descripcion_grafico = ""
        numero_fila_elegido = fila_datos.get('Numero', 'N/A')
                                            
        classification_details = {
            "ID": fila_datos.get("ID"),
            "GRADO": grado_elegido,
            "ÁREA": area_elegida,
            "ASIGNATURA": asignatura_elegida,
            "MACROHABILIDAD": macrohabilidad_elegida,
            "PROCESO COGNITIVO": proceso_cognitivo_elegido,
            "MICROHABILIDAD": microhabilidad_elegida,
            "COMPETENCIA MICROHABILIDAD": competencia_microhabilidad_elegida,
            "Numero": numero_fila_elegido
        }
    
        item_final_data = None
        full_generation_prompt = ""
        full_auditor_prompt = ""
    
        if feedback_usuario and item_a_refinar_text:
            prompt_refinamiento = f"""
            --- TAREA DE REFINAMIENTO ---
            Eres un experto en ítems de evaluación. Tu tarea es REFINAR el siguiente ítem, corrigiendo o ajustando el texto para abordar las observaciones del usuario. No lo reescribas completamente; solo haz las correcciones necesarias.
            
            --- OBSERVACIONES DEL USUARIO PARA REFINAR ---
            {feedback_usuario}
            -------------------------------------------
            
            --- ÍTEM ORIGINAL A REFINAR ---
            {item_a_refinar_text}
            -----------------------------
            
            --- INSTRUCCIONES DE SALIDA ---
            Devuelve el ítem refinado con el mismo formato original:
            PREGUNTA: ...
            A. ...
            B. ...
            C. ...
            D. ...
            RESPUESTA CORRECTA: ...
            JUSTIFICACIONES:
            A. ...
            B. ...
            C. ...
            D. ...
            GRAFICO_NECESARIO: [SÍ/NO]
            DESCRIPCION_GRAFICO: [Descripción detallada o N/A]
            """
            
            try:
                # Quitamos el spinner detallado de aquí
                full_llm_response = generar_texto_con_llm(gen_model_name, prompt_refinamiento)
            
                if full_llm_response is None:
                    st.error("Fallo en la refinación del ítem.")
                    return None
            
                item_and_graphic_match = re.search(r"(PREGUNTA:.*?)(GRAFICO_NECESARIO:\s*(SÍ|NO).*?DESCRIPCION_GRAFICO:.*)", full_llm_response, re.DOTALL)
                if item_and_graphic_match:
                    current_item_text = item_and_graphic_match.group(1).strip()
                    grafico_info_block = item_and_graphic_match.group(2).strip()
                    grafico_necesario_match = re.search(r"GRAFICO_NECESARIO:\s*(SÍ|NO)", grafico_info_block)
                    if grafico_necesario_match:
                        grafico_necesario = grafico_necesario_match.group(1).strip()
                    descripcion_grafico_match = re.search(r"DESCRIPCION_GRAFICO:\s*(.*)", grafico_info_block, re.DOTALL)
                    if descripcion_grafico_match:
                        # --- INICIO DEL NUEVO BLOQUE DE CÓDIGO ---
                        descripcion_grafico_str = descripcion_grafico_match.group(1).strip()
                        descripciones_graficos_list = [] # Nueva variable para la lista

                        if descripcion_grafico_str.upper() != 'N/A' and descripcion_grafico_str.strip().startswith('['):
                            try:
                                descripciones_graficos_list = json.loads(descripcion_grafico_str)
                            except json.JSONDecodeError:
                                print("Error al decodificar JSON de gráficos (generación). Tratando como texto simple.")
                                descripciones_graficos_list = [{"ubicacion": "enunciado", "tipo_elemento": "otro_tipo", "datos": {"descripcion_natural": descripcion_grafico_str}}]
                        elif descripcion_grafico_str.upper() != 'N/A':
                             # Si no es una lista, lo tratamos como un solo gráfico para enunciado
                            descripciones_graficos_list = [{"ubicacion": "enunciado", "tipo_elemento": "otro_tipo", "datos": {"descripcion_natural": descripcion_grafico_str}}]

                        # Asignamos la lista a la variable original para que el resto del código la use
                        descripcion_grafico = descripciones_graficos_list
                        # --- FIN DEL NUEVO BLOQUE DE CÓDIGO ---
                else:
                    current_item_text = full_llm_response
                    grafico_necesario = "NO"
                    descripcion_grafico = ""
            
                auditoria_resultado, full_auditor_prompt = auditar_item_con_llm(
                    audit_model_name,
                    item_generado=current_item_text,
                    grado=grado_elegido, area=area_elegida, asignatura=asignatura_elegida, macrohabilidad=macrohabilidad_elegida,
                    proceso_cognitivo=proceso_cognitivo_elegido,
                    microhabilidad=microhabilidad_elegida, competencia_microhabilidad=competencia_microhabilidad_elegida,
                    contexto_educativo=contexto_educativo, manual_reglas_texto=manual_reglas_texto,
                    descripcion_bloom=descripcion_bloom,
                    grafico_necesario=grafico_necesario,
                    descripcion_grafico=descripcion_grafico,
                    prompt_auditor_adicional=prompt_auditor_adicional
                )
    
                dictamen_final_match = re.search(r"DICTAMEN FINAL:\s*\[(.*?)]", auditoria_resultado, re.DOTALL)
                auditoria_status = dictamen_final_match.group(1).strip() if dictamen_final_match else "❌ RECHAZADO (no se pudo extraer dictamen)"
                observaciones_start = auditoria_resultado.find("OBSERVACIONES FINALES:")
                audit_observations = auditoria_resultado[observaciones_start + len("OBSERVACIONES FINALES:"):].strip() if observaciones_start != -1 else "No se pudieron extraer observaciones específicas."
                
                item_final_data = {
                    "item_text": current_item_text,
                    "classification": classification_details,
                    "grafico_necesario": grafico_necesario,
                    "descripciones_graficos": descripcion_grafico,
                    "final_audit_status": auditoria_status,
                    "final_audit_observations": audit_observations,
                    "generation_prompt_used": prompt_refinamiento,
                    "auditor_prompt_used": full_auditor_prompt
                }
                return item_final_data
            
            except Exception as e:
                audit_observations = f"Error técnico durante la refinación: {e}. Por favor, corrige este problema."
                auditoria_status = "❌ RECHAZADO (error técnico)"
                item_final_data = {
                    "item_text": current_item_text if current_item_text else "No se pudo refinar el ítem debido a un error técnico.",
                    "classification": classification_details,
                    "grafico_necesario": "NO",
                    "descripcion_grafico": "",
                    "descripciones_graficos": "",
                    "final_audit_status": auditoria_status,
                    "final_audit_observations": audit_observations,
                    "generation_prompt_used": prompt_refinamiento,
                    "auditor_prompt_used": full_auditor_prompt
                }
                return item_final_data
        else:
            while auditoria_status != "✅ CUMPLE TOTALMENTE" and attempt < max_refinement_attempts:
                attempt += 1
                
                # --- INICIO DE LA NUEVA LÓGICA REFORZADA PARA EL CONTEXTO ---
                instruccion_contexto = ""
                formato_salida_pregunta = "PREGUNTA: [Redacta aquí el contexto (si es necesario) y el enunciado de la pregunta]"
    
                if contexto_general_macrohabilidad:
                    # Si SÍ hay un contexto compartido, las instrucciones son estrictas.
                    instruccion_contexto = f"""
    --- CONTEXTO GENERAL OBLIGATORIO DE LA MACROHABILIDAD ---
    ¡INSTRUCCIÓN CRÍTICA! Debes iniciar el campo 'PREGUNTA:' exactamente con el siguiente texto de contexto, sin alterarlo, resumirlo o parafrasearlo. Después del contexto, redacta el enunciado específico para el ítem.
    
    CONTEXTO GENERAL DE LA MACROHABILIDAD (DEBE SER INCLUIDO TEXTUALMENTE):
    "{contexto_general_macrohabilidad}"
    ----------------------------------------------------
    """
                    formato_salida_pregunta = "PREGUNTA: [Texto del CONTEXTO GENERAL DE LA MACROHABILIDAD] [Enunciado específico de la pregunta]"
                else:
                    # Si NO hay contexto, las instrucciones son las normales.
                    instruccion_contexto = """
    --- CONTEXTO GENERAL DE LA MACROHABILIDAD (si aplica) ---
    Este ítem debe generar su propio contexto individual, ya que no se ha definido un contexto general para la macrohabilidad.
    ----------------------------------------------------
    """
                # --- FIN DE LA NUEVA LÓGICA ---
    
                clave_aleatoria = random.choice(['A', 'B', 'C', 'D'])

                seccion_contexto_libro = ""
                if contexto_del_libro:
                    seccion_contexto_libro = f"""
        --- CONTEXTO PRINCIPAL DEL LIBRO GUÍA (¡USO OBLIGATORIO!) ---
        ¡INSTRUCCIÓN CRÍTICA! Debes basar tu pregunta, opciones y justificaciones **principalmente** en los siguientes extractos del libro guía. La respuesta correcta DEBE poder deducirse de este texto.
        
        {contexto_del_libro}
        ---------------------------------------------------------
        """
                #
                seccion_imagen = ""
                if descripcion_imagen_aprobada:
                    seccion_imagen = f"""
    --- INFORMACIÓN VISUAL OBLIGATORIA (BASADA EN IMAGEN) ---
    ¡INSTRUCCIÓN CRÍTICA! El ítem que construyas DEBE basarse directamente en la siguiente descripción de una imagen. La pregunta, las opciones y las justificaciones deben hacer referencia a los detalles mencionados aquí. Este es el insumo principal.
    
    DESCRIPCIÓN DE LA IMAGEN:
    "{descripcion_imagen_aprobada}"
    ---------------------------------------------------------
    """
    
                prompt_content_for_llm = f"""
                Eres un psicómetra y diseñador experto en ítems de evaluación educativa, con profundo conocimiento en la Taxonomía de Bloom y su aplicación práctica.
                Tu tarea es construir un ítem de {tipo_pregunta} con una única respuesta correcta, garantizando una alineación perfecta y demostrable con el marco cognitivo solicitado, siguiendo un riguroso proceso de análisis previo.
                
                --- CONTEXTO Y PARÁMETROS DEL ÍTEM ---
                - Grado: {grado_elegido}
                - Área: {area_elegida}
                - Asignatura: {asignatura_elegida}
                - Macrohabilidad o unidad temática: {macrohabilidad_elegida}
                - Proceso cognitivo (Taxonomía de Bloom): {proceso_cognitivo_elegido}
                - Descripción DETALLADA y VINCULANTE del proceso cognitivo:
                    "{descripcion_bloom}"

                --- EJEMPLOS Y GUÍAS DE PREGUNTAS (Preguntas Tipo) ---
                ¡INSTRUCCIÓN CLAVE! Para asegurar que el enunciado del ítem se alinee con el proceso cognitivo, inspírate en los siguientes ejemplos. La pregunta que formules debe seguir un estilo similar, buscando una única respuesta correcta y evitando comparaciones subjetivas ("mejor", "más adecuado").
                {ejemplos_preguntas}
                ----------------------------------------------------
                
                --- PROMPT ADICIONAL: TAXONOMÍA DE BLOOM / PROCESOS COGNITIVOS ---
                {prompt_bloom_adicional if prompt_bloom_adicional else "No se proporcionaron prompts adicionales específicos para taxonomía de Bloom."}
                ------------------------------------------------------------------
                
                - Microhabilidad (foco principal del ítem): {microhabilidad_elegida}
                - Nivel educativo esperado del estudiante: {contexto_educativo}
                - Nivel de dificultad deseado: {dificultad}
                
                {instruccion_contexto}

                {seccion_imagen}

                {seccion_contexto_libro}

                # =============================================================================
                # INICIO DE LA MODIFICACIÓN CLAVE: ANÁLISIS COGNITIVO OBLIGATORIO Y EXCLUSIVO
                # =============================================================================
                --- ANÁLISIS COGNITIVO OBLIGATORIO (TAXONOMÍA DE BLOOM) ---
                Antes de escribir el ítem, DEBES realizar el siguiente análisis interno para garantizar una alineación perfecta. La calidad de tu pregunta dependerá de la rigurosidad de este análisis.
                
                1.  **Deconstrucción del Proceso Cognitivo**: Revisa la "Descripción DETALLADA y VINCULANTE del proceso cognitivo" proporcionada. Es de carácter **obligatorio** que extraigas de ella el subproceso y los **verbos de acción clave** o sinónimos directos que mejor se alineen con la microhabilidad '{microhabilidad_elegida}'.
                
                2.  **Diseño de la Tarea Cognitiva**: Describe la tarea mental específica y observable que el estudiante DEBE realizar. **Esta descripción debe incorporar explícitamente los verbos de acción (o sus sinónimos directos) que identificaste en el paso anterior.** No describas la pregunta, sino la operación mental. (Ej: "La tarea exige que el estudiante *compare* dos eventos históricos para *detectar correspondencias* entre sus causas económicas, y luego *construya un modelo* simple de causa-efecto que *explique* esas similitudes.").
                
                3.  **Justificación de la Alineación**: Justifica explícitamente cómo la "Tarea Cognitiva" que diseñaste se alinea con la definición del proceso "{proceso_cognitivo_elegido}" y su subproceso. (Ej: "Esta tarea se alinea con COMPRENDER-Comparar y Explicar porque el estudiante debe procesar información, detectar relaciones y construir un modelo causal, lo cual va más allá de solo recordar los hechos.").
                
                4.  **Verificación de Exclusividad Cognitiva (¡CRÍTICO!)**: Debes confirmar que la tarea diseñada NO pertenece a otros niveles cognitivos. Justifica brevemente por qué la tarea:
                    * **Supera el nivel anterior**: Explica por qué la tarea es más compleja que el nivel cognitivo inmediatamente inferior en la taxonomía. (Ej: "No es solo RECORDAR porque no se pide evocar fechas, sino relacionarlas.").
                    * **No alcanza el nivel superior**: Explica por qué la tarea no llega a la complejidad del nivel cognitivo inmediatamente superior. (Ej: "No es ANALIZAR porque no se le pide que deconstruya la validez de las fuentes de información o que determine sesgos, solo que organice y explique la información presentada.").
                
                La pregunta que construirás a continuación debe ser la materialización exacta de esta Tarea Cognitiva verificada.
                # =============================================================================
                # FIN DE LA MODIFICACIÓN CLAVE
                # =============================================================================
                
                --- INSTRUCCIONES PARA LA CONSTRUCCIÓN DEL ÍTEM ---
                CONTEXTO DEL ÍTEM:
                - Debe ser relevante y plausible, sirviendo como el escenario donde se ejecutará la Tarea Cognitiva que diseñaste.
                - La temática debe ser la de la {macrohabilidad_elegida} y ser central para el problema.
                - Evita referencias a marcas, nombres propios, lugares reales o información personal identificable.
                
                ENUNCIADO:
                - **CRÍTICO**: Formula una pregunta que fuerce al estudiante a ejecutar la Tarea Cognitiva que definiste y verificaste en tu análisis. El enunciado es el disparador de esa operación mental.
                - Formula una pregunta clara, directa, sin ambigüedades ni tecnicismos innecesarios.
                - ¡INSTRUCCIÓN CRÍTICA DE ESTILO! Evita terminantemente formular preguntas que pidan al estudiante comparar o jerarquizar opciones. **NO USES** frases como "¿cuál es la opción más...", "¿cuál es el mejor...", "¿cuál describe principalmente...?", "¿cuál es la razón principal...?". La pregunta debe tener una única respuesta objetivamente correcta.
                - En su lugar, formula preguntas directas como: "**¿Cuál es la causa de...?**", "**¿Qué conclusión se deriva de...?**", "**¿Cuál de las afirmaciones es correcta?**".
                - NO uses directamente en la pregunta el verbo principal del proceso cognitivo (ej. no preguntes "¿Cuál es el análisis de...?"). Busca redacciones más auténticas.
                - Si utilizas negaciones, resáltalas en MAYÚSCULAS Y NEGRITA (por ejemplo: **NO ES**, **EXCEPTO**).
                
                OPCIONES DE RESPUESTA:
                - Escribe exactamente cuatro opciones (A, B, C y D).
                - **Opción Correcta**: Debe ser la única conclusión válida tras ejecutar correctamente la Tarea Cognitiva.
                - La respuesta correcta DEBE ser la opción {clave_aleatoria}.
                - **Distractores (Incorrectos)**: Deben ser plausibles y diseñados a partir de errores típicos en la ejecución de la Tarea Cognitiva. (Ej: un distractor podría ser el resultado de aplicar un proceso cognitivo inferior, como simplemente recordar un dato, en lugar de analizarlo).
                - Las respuestas deben tener una estructura gramatical y longitud similares.
                - No utilices fórmulas vagas como “ninguna de las anteriores” o “todas las anteriores”.
                
                JUSTIFICACIONES:
                {formato_justificacion}
    
                --- PROMPT ADICIONAL: REGLAS GENERALES DE CONSTRUCCIÓN ---
                {prompt_construccion_adicional if prompt_construccion_adicional else "No se proporcionaron prompts adicionales específicos para reglas generales de construcción."}
                ---------------------------------------------------------
    
                --- REGLAS ADICIONALES DEL MANUAL DE CONSTRUCCIÓN ---
                Considera y aplica estrictamente todas las directrices, ejemplos y restricciones contenidas en el siguiente manual.
                Esto es de suma importancia para la calidad y pertinencia del ítem.
    
                Manual de Reglas:
                {manual_reglas_texto}
                ----------------------------------------------------
    
                --- INFORMACIÓN ADICIONAL PROPORCIONADA POR EL USUARIO (Contexto General) ---
                {informacion_adicional_usuario if informacion_adicional_usuario else "No se proporcionó información adicional general."}
                ---------------------------------------------------------------------------
                
                --- PROMPT ADICIONAL: COSAS ESPECÍFICAS A TENER EN CUENTA ---
                {prompt_especifico_adicional if prompt_especifico_adicional else "No se proporcionaron prompts adicionales específicos para consideraciones adicionales."}
                ----------------------------------------------------------
    
                --- DATO CLAVE PARA LA CONSTRUCCIÓN ---
                Basado en el foco temático y el proceso cognitivo, considera el siguiente dato o idea esencial:
                "{dato_para_pregunta_foco}"
    
                --- INSTRUCCIONES ESPECÍFICAS DE SALIDA PARA GRÁFICO ---
                Después del bloque de JUSTIFICACIONES, indica si el ítem necesita elementos visuales.
                
                ¡INSTRUCCIÓN CRÍTICA! **Considera como elemento visual cualquier cosa que no sea texto de prosa**, incluyendo: gráficos, diagramas, **tablas**, construcciones geométricas, etc.
                
                GRAFICO_NECESARIO: [SÍ/NO]
                DESCRIPCION_GRAFICO: [Si la respuesta es **NO**, escribe **N/A**. Si la respuesta es **SÍ**, DEBES proporcionar una **LISTA DE OBJETOS JSON VÁLIDOS** incluso si solo hay un gráfico, siguiendo estas reglas:]
                El JSON siempre debe contener los campos: `"ubicacion"`, `"tipo_elemento"`, `"datos"`, `"configuracion"` y `"descripcion"`.

                1. Cada objeto DEBE contener una clave `"ubicacion"` para identificar dónde va el gráfico. Usa uno de los siguientes valores: `"enunciado"`, `"opcion_a"`, `"opcion_b"`, `"opcion_c"`, `"opcion_d"`.
                
                2.  Para `"tipo_elemento"`, elige **UNO** de la siguiente lista: `grafico_barras_verticales`, `grafico_circular`, `tabla`, `construccion_geometrica`, `diagrama_arbol`, `flujograma`, `pictograma`, `scatter_plot`, `line_plot`, `histogram`, `box_plot`, `violin_plot`, `heatmap`, `contour_plot`, `3d_plot`, `network_diagram`, `area_plot`, `radar_chart`, `venn_diagram`, `fractal`, `otro_tipo`.
                
                3.  Para `"descripcion"`, proporciona un **texto en lenguaje natural que resuma y detalle todos los elementos clave del gráfico**, sus relaciones y las características que se deben tener en cuenta para generarlo visualmente.
                
                4.  **LÓGICA CONDICIONAL PARA EL CAMPO "datos":**
                    * **Si eliges un `tipo_elemento` de la lista (QUE NO SEA `otro_tipo`)**: El campo `"datos"` debe ser un objeto con la **información estructurada y numérica**.
                        * *Ejemplo para `tabla`*:
                        ```json
                        {{
                          "ubicacion": "enunciado",
                          "tipo_elemento": "tabla",
                          "datos": {{
                            "columnas": ["País", "Capital"],
                            "filas": [["Colombia", "Bogotá"], ["Argentina", "Buenos Aires"]]
                          }},
                          "configuracion": {{ "titulo": "Capitales de Sudamérica" }},
                          "descripcion": "Una tabla de dos columnas que lista países sudamericanos y sus respectivas capitales. La primera columna corresponde al país y la segunda a su capital."
                        }}
                        ```
                    * **Si el gráfico no corresponde a ninguno y eliges `otro_tipo`**: El campo `"datos"` debe contener un único objeto con la clave `"descripcion_natural"`, cuyo valor será un **texto exhaustivo** con todos los detalles necesarios para construir el gráfico desde cero.
                        * *Ejemplo para `otro_tipo`*:
                        ```json
                        {{
                          "ubicacion": "opcion_a",
                          "tipo_elemento": "otro_tipo",
                          "datos": {{
                            "descripcion_natural": "Se requiere un diagrama de un circuito eléctrico simple en serie. Debe mostrar una fuente de poder (batería) de 9V conectada a tres resistencias (R1=10Ω, R2=20Ω, R3=30Ω) una después de la otra. El diagrama debe indicar claramente la dirección del flujo de la corriente (I) con una flecha saliendo del polo positivo de la batería."
                          }},
                          "configuracion": {{ "titulo": "Circuito en Serie" }},
                          "descripcion": "Diagrama de un circuito eléctrico simple con una batería y tres resistencias conectadas en serie, mostrando el flujo de la corriente."
                        }}
                        ```
 
                --- FORMATO ESPERADO DE SALIDA ---
                ¡INSTRUCCIÓN CRÍTICA! Tu respuesta DEBE ser un único bloque de código JSON válido, sin ningún otro texto o explicación antes o después (no uses \`\`\`json).
                El objeto JSON debe tener la siguiente estructura:              
                {{
                  "pregunta": "Aquí va el texto del contexto (si lo hay) seguido del enunciado de la pregunta.",
                  "opciones": {{
                    "A": "Texto de la opción A.",
                    "B": "Texto de la opción B.",
                    "C": "Texto de la opción C.",
                    "D": "Texto de la opción D."
                  }},
                  "respuestaCorrecta": "{clave_aleatoria}",
                  "justificaciones": {{
                    "A": "Justificación para la opción A.",
                    "B": "Justificación para la opción B.",
                    "C": "Justificación para la opción C.",
                    "D": "Justificación para la opción D."
                  }},
                  "graficoNecesario": "SÍ",
                  "descripcionGrafico": [
                    {{
                      "ubicacion": "enunciado",
                      "tipo_elemento": "tabla",
                      "datos": {{"columnas": ["X"], "filas": [[1]]}},
                      "configuracion": {{"titulo": "Ejemplo"}},
                      "descripcion": "Descripción del gráfico."
                    }}
                  ]
                }}

                Asegúrate de que el valor de "respuestaCorrecta" sea exactamente "{clave_aleatoria}". Si "graficoNecesario" es "NO", el valor de "descripcionGrafico" debe ser un array vacío [].
                """
                
                if attempt > 1:
                    prompt_content_for_llm += f"""
                    --- RETROALIMENTACIÓN DE AUDITORÍA PARA REFINAMIENTO ---
                    El ítem anterior no cumplió con todos los criterios. Por favor, revisa las siguientes observaciones y mejora el ítem para abordarlas.
                    Observaciones del Auditor:
                    {audit_observations}
                    ---------------------------------------------------
                    --- ÍTEM ANTERIOR A REFINAR ---
                    {current_item_text}
                    -------------------------------
                    """
                
                full_generation_prompt = prompt_content_for_llm
    
                try:
                    # Quitamos el spinner de aquí
                    full_llm_response = generar_texto_con_llm(gen_model_name, prompt_content_for_llm)
                    
                    if full_llm_response is None:
                        auditoria_status = "❌ RECHAZADO (Error de Generación)"
                        audit_observations = "El modelo de generación no pudo producir una respuesta válida."
                        break
                    
                    # ... (dentro del while y el primer try/except) ...
                    # -- INICIO DEL NUEVO BLOQUE DE PARSEO JSON --
                    try:
                        # --- INICIO DE LA MODIFICACIÓN ---
                        # 1. Busca el inicio del JSON (el primer '{') y el final (el último '}')
                        json_start = full_llm_response.find('{')
                        json_end = full_llm_response.rfind('}')

                        # 2. Verifica si se encontró un objeto JSON en la respuesta
                        if json_start != -1 and json_end != -1:
                            # 3. Extrae únicamente el string del JSON
                            json_string = full_llm_response[json_start:json_end + 1]
                            # 4. Intenta decodificar SÓLO el string extraído
                            item_data = json.loads(json_string)
                        else:
                            # Si no se encontró, lanza un error para que sea manejado por el bloque 'except'
                            raise json.JSONDecodeError("No se encontró un objeto JSON en la respuesta del modelo.", full_llm_response, 0)
                        # --- FIN DE LA MODIFICACIÓN ---

                        # 2. Reconstruimos el texto del ítem para mostrarlo en la UI y para el auditor
                        opciones_texto = "\n".join([f"{key}. {value}" for key, value in item_data.get("opciones", {}).items()])
                        justificaciones_texto = "\n".join([f"{key}. {value}" for key, value in item_data.get("justificaciones", {}).items()])
                        
                        current_item_text = (
                            f"PREGUNTA: {item_data.get('pregunta', '')}\n"
                            f"{opciones_texto}\n"
                            f"RESPUESTA CORRECTA: {item_data.get('respuestaCorrecta', '')}\n"
                            f"JUSTIFICACIONES:\n{justificaciones_texto}"
                        )
                        
                        # 3. Extraemos la información del gráfico directamente
                        grafico_necesario = item_data.get("graficoNecesario", "NO")
                        # El prompt ya pide que sea una lista de objetos, así que la obtenemos directamente
                        descripciones_graficos_list = item_data.get("descripcionGrafico", [])
                        descripcion_grafico = descripciones_graficos_list # Asignamos para la auditoría


                    except json.JSONDecodeError:
                        # Si el LLM no devuelve un JSON válido, lo marcamos como un error de formato
                        auditoria_status = "❌ RECHAZADO (Error de Formato JSON)"
                        audit_observations = f"El modelo de generación no produjo un JSON válido. Salida recibida:\n{full_llm_response}"
                        st.warning(audit_observations)
                        current_item_text = full_llm_response # Guardamos el texto erróneo para el reintento
                        grafico_necesario = "NO"
                        descripcion_grafico = ""
                        # Forzamos la salida del bucle de reintentos si hay error de formato
                    # -- FIN DEL NUEVO BLOQUE DE PARSEO JSON --
                    
                    # Quitamos el spinner de aquí

                    auditoria_json_str, full_auditor_prompt = auditar_item_con_llm(
                        audit_model_name,
                        item_generado=current_item_text,
                        grado=grado_elegido, area=area_elegida, asignatura=asignatura_elegida, macrohabilidad=macrohabilidad_elegida,
                        proceso_cognitivo=proceso_cognitivo_elegido,
                        microhabilidad=microhabilidad_elegida, competencia_microhabilidad=competencia_microhabilidad_elegida,
                        contexto_educativo=contexto_educativo, manual_reglas_texto=manual_reglas_texto,
                        descripcion_bloom=descripcion_bloom,
                        grafico_necesario=grafico_necesario,
                        descripcion_grafico=descripcion_grafico,
                        prompt_auditor_adicional=prompt_auditor_adicional
                    )
    
                    if auditoria_json_str is None:
                        auditoria_status = "❌ RECHAZADO (Error de Auditoría)"
                        audit_observations = "El modelo de auditoría no pudo producir una respuesta válida."
                    else:
                        auditoria_data = parse_json_llm(auditoria_json_str)
                        if auditoria_data is None:
                            auditoria_status = "❌ RECHAZADO (Error de Formato JSON del Auditor)"
                            audit_observations = f"El modelo de auditoría no produjo un JSON válido. Salida recibida:\n{auditoria_json_str}"
                            st.warning(audit_observations)
                        else:
                            auditoria_status = auditoria_data.get("dictamen_final", "❌ RECHAZADO (Clave no encontrada)")
                            audit_observations = auditoria_data.get("observaciones_finales", "No se encontraron observaciones.")
                    
                    item_final_data = {
                        "item_text": current_item_text,
                        "classification": classification_details,
                        "grafico_necesario": grafico_necesario,
                        "descripciones_graficos": descripciones_graficos_list,
                        "final_audit_status": auditoria_status,
                        "final_audit_observations": audit_observations,
                        "generation_prompt_used": full_generation_prompt,
                        "auditor_prompt_used": full_auditor_prompt
                    }
        
                    if auditoria_status == "✅ CUMPLE TOTALMENTE":
                        break

                
                except Exception as e:
                    audit_observations = f"Error técnico durante la generación: {e}. Por favor, corrige este problema."
                    auditoria_status = "❌ RECHAZADO (error técnico)"
                    item_final_data = {
                        "item_text": current_item_text if current_item_text else "No se pudo generar el ítem debido a un error técnico.",
                        "classification": classification_details,
                        "grafico_necesario": "NO",
                        "descripcion_grafico": "",
                        "final_audit_status": auditoria_status,
                        "final_audit_observations": audit_observations,
                        "generation_prompt_used": full_generation_prompt,
                        "auditor_prompt_used": full_auditor_prompt
                    }
                    break
    
            if item_final_data is None:
                return None
    
            return item_final_data

    def crear_documento_word_individual(item_data):
        """
        Crea un documento de Word en memoria para un único ítem procesado.
        """
        doc = docx.Document()
    
        # Extraer datos del ítem
        pregunta_texto = item_data.get("item_text", "No disponible")
        classification = item_data.get("classification", {})
    
        # --- AÑADIR NÚMERO DE FILA DESDE LA COLUMNA "Numero" ---
        numero_de_fila = classification.get("Numero", "N/A")
        p_fila = doc.add_paragraph()
        run_fila_label = p_fila.add_run("Fila: ")
        run_fila_label.bold = True
        p_fila.add_run(str(numero_de_fila))
        doc.add_paragraph('')  # Espacio después
    
        # Añadir clasificación
        doc.add_paragraph('--- Clasificación del Ítem ---')
        for key, value in classification.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{key}: ")
            run.bold = True
            p.add_run(str(value))
    
        doc.add_paragraph('')
    
        # --- INICIO DE LA MODIFICACIÓN ---
        # Añadir la imagen de origen si existe
        if 'source_image' in item_data and item_data['source_image']:
            doc.add_heading('Insumo Visual de Origen', level=2)
            try:
                # Regresamos al inicio del buffer de la imagen
                item_data['source_image'].seek(0)
                doc.add_picture(item_data['source_image'], width=Inches(5.0))
            except Exception as e:
                doc.add_paragraph(f"No se pudo incrustar la imagen de origen. Error: {e}")
            doc.add_paragraph('')
        # --- FIN DE LA MODIFICACIÓN ---

        # Añadir texto del ítem (pregunta, opciones, etc.)
        doc.add_paragraph('--- Contenido del Ítem ---')
        doc.add_paragraph(pregunta_texto)
        doc.add_paragraph('')
    
        # --- Obtener info de gráfico del item_data (evita NameError) ---
        grafico_necesario_val = str(item_data.get("grafico_necesario", "NO") or "NO")
        desc_grafico = item_data.get("descripcion_grafico", "") or ""
    
        # --- Gráfico (solo si es requerido) ---
        def _es_si(s):
            s = str(s).lower().strip().replace("í", "i")  # normaliza "sí" -> "si"
            return s == "si"
               
        # --- INICIO DEL BLOQUE DE ANEXO FINAL ---
        grafico_necesario_val = str(item_data.get("grafico_necesario", "NO") or "NO").lower().strip()
    
        if grafico_necesario_val == 'sí' or grafico_necesario_val == 'si':
            descripciones_graficos = item_data.get("descripciones_graficos", [])
            imagenes_guardadas = item_data.get("generated_images", [])
            
            if descripciones_graficos:
                doc.add_page_break()
                doc.add_heading('Anexo de Gráficos', level=1)
                
                for i, desc_grafico in enumerate(descripciones_graficos):
                    ubicacion = desc_grafico.get("ubicacion", f"gráfico_{i+1}")
                    ubicacion_titulo = ubicacion.replace("_", " ").title()
                    
                    doc.add_heading(f"Gráfico {i+1}: Para {ubicacion_titulo}", level=2)
                    
                    # Buscamos la imagen correspondiente a esta descripción
                    imagen_encontrada = None
                    for img in imagenes_guardadas:
                        if img.get("ubicacion") == ubicacion:
                            imagen_encontrada = img.get("buffer")
                            break
                    
                    if imagen_encontrada:
                        # ¡Incrustamos la imagen!
                        imagen_encontrada.seek(0) # Regresamos al inicio del buffer de la imagen
                        doc.add_picture(imagen_encontrada, width=Inches(5.5))
                    else:
                        # Si no hay imagen, guardamos la descripción como respaldo
                        doc.add_paragraph("No se generó una imagen para este gráfico. Se adjunta su descripción técnica:")
                        json_str = json.dumps(desc_grafico, indent=2, ensure_ascii=False)
                        p = doc.add_paragraph(json_str)
                        p.style = 'Quote'
        # --- FIN DEL BLOQUE DE ANEXO FINAL ---
    
        # Guardar en un buffer de memoria y devolverlo
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def exportar_a_zip(preguntas_procesadas_list):
        """
        Crea un archivo .zip en memoria que contiene un .docx individual para cada ítem aprobado.
        Nombra cada .docx según su ID y maneja duplicados con sufijos (_A, _B, C...).
        """
        zip_buffer = io.BytesIO()
        id_counts = {}

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for item_data in preguntas_procesadas_list:
                item_id = item_data.get("classification", {}).get("ID", "sin_id")
                
                # Contar cuántas veces hemos usado este ID
                count = id_counts.get(item_id, 0)
                
                # Generar el nombre de archivo único
                if count > 0:
                    # Añadir sufijo _B, _C, etc. (chr(65) es 'A')
                    sufijo = f"_{chr(65 + count)}"
                    nombre_archivo = f"{item_id}{sufijo}.docx"
                else:
                    # Para la primera aparición, podría ser _A o sin sufijo. Usemos _A por consistencia.
                    nombre_archivo = f"{item_id}_A.docx"

                id_counts[item_id] = count + 1

                # Crear el documento Word individual en memoria
                doc_buffer = crear_documento_word_individual(item_data)
                
                # Añadir el archivo Word al ZIP
                zip_file.writestr(nombre_archivo, doc_buffer.getvalue())

        zip_buffer.seek(0)
        return zip_buffer
        
    # --- NUEVA FUNCIÓN: Generar imagen con IA (Indentación y robustez corregidas) ---
    def generar_imagen_con_ia(prompt_descripcion: str):
        """
        Genera una imagen desde una descripción textual usando Vertex AI.
        Devuelve un BytesIO con la imagen o None si falla.
        """
        if not prompt_descripcion or not str(prompt_descripcion).strip():
            st.warning("El prompt de la imagen está vacío.")
            return None
    
        try:
            # Asegúrate de haber llamado vertexai.init(project=..., location=...) antes en tu app.
            modelo_imagen = ImageGenerationModel.from_pretrained("imagen-4.0-fast-generate-001")
    
            respuesta = modelo_imagen.generate_images(
                prompt=prompt_descripcion,
                number_of_images=1,
                # Ejemplos de parámetros opcionales:
                # aspect_ratio="1:1",
                # negative_prompt="texto borroso, baja resolución",
            )
    
            if not respuesta or not getattr(respuesta, "images", None):
                raise ValueError("El modelo no devolvió imágenes.")
    
            img_obj = respuesta.images[0]
    
            # Preferimos atributo público si existe; si no, fallback al interno (_image_bytes)
            image_bytes = getattr(img_obj, "image_bytes", None) or getattr(img_obj, "_image_bytes", None)
    
            # Algunos SDKs exponen método para bytes
            if image_bytes is None and hasattr(img_obj, "as_bytes"):
                image_bytes = img_obj.as_bytes()
    
            if image_bytes is None:
                raise ValueError("No fue posible extraer los bytes de la imagen.")
    
            return io.BytesIO(image_bytes)
    
        except Exception as e:
            st.error(f"Error al generar imagen con IA: {e}")
            return None
    # --- FIN DE LA NUEVA FUNCIÓN ---
         
    def exportar_a_excel(preguntas_procesadas_list, nombre_archivo_base):
        """
        Exporta los ítems procesados a un archivo Excel con una estructura detallada.
        """
        datos_para_excel = []
    
        for i, item_data in enumerate(preguntas_procesadas_list):
            item_text = item_data.get("item_text", "")
            classification = item_data.get("classification", {})
        
            pregunta_bloque_match = re.search(r"PREGUNTA:(.*?)(?=A\.)", item_text, re.S)
            contexto = ""
            enunciado = ""
            if pregunta_bloque_match:
                bloque_completo = pregunta_bloque_match.group(1).strip()
                ultimo_interrogante = bloque_completo.rfind('?')
                if ultimo_interrogante != -1:
                    contexto = bloque_completo[:ultimo_interrogante+1].strip()
                    enunciado = bloque_completo[ultimo_interrogante+1:].strip()
                else:
                    contexto = bloque_completo
            
            opcion_a_match = re.search(r"\nA\.\s(.*?)(?=\nB\.)", item_text, re.S)
            opcion_b_match = re.search(r"\nB\.\s(.*?)(?=\nC\.)", item_text, re.S)
            opcion_c_match = re.search(r"\nC\.\s(.*?)(?=\nD\.)", item_text, re.S)
            opcion_d_match = re.search(r"\nD\.\s(.*?)(?=\nRESPUESTA CORRECTA:)", item_text, re.S)
            
            clave_match = re.search(r"RESPUESTA CORRECTA:\s*(\w)", item_text)
            
            justificaciones_bloque_match = re.search(r"JUSTIFICACIONES:(.*)", item_text, re.S)
            just_a, just_b, just_c, just_d = "", "", "", ""
            if justificaciones_bloque_match:
                just_texto = justificaciones_bloque_match.group(1)
                just_a_match = re.search(r"A\.\s(.*?)(?=\n\s*B\.|\Z)", just_texto, re.S)
                just_b_match = re.search(r"B\.\s(.*?)(?=\n\s*C\.|\Z)", just_texto, re.S)
                just_c_match = re.search(r"C\.\s(.*?)(?=\n\s*D\.|\Z)", just_texto, re.S)
                just_d_match = re.search(r"D\.\s(.*)", just_texto, re.S)
                if just_a_match: just_a = just_a_match.group(1).strip()
                if just_b_match: just_b = just_b_match.group(1).strip()
                if just_c_match: just_c = just_c_match.group(1).strip()
                if just_d_match: just_d = just_d_match.group(1).strip()
                
            fila = {
                'item': f"{nombre_archivo_base}_{i+1}",
                'ID': classification.get('ID'),
                'GRADO': classification.get('GRADO'),
                'ÁREA': classification.get('ÁREA'),
                'ASIGNATURA': classification.get('ASIGNATURA'),
                'MACROHABILIDAD': classification.get('MACROHABILIDAD'),
                'PROCESO COGNITIVO': classification.get('PROCESO COGNITIVO'),
                'MICROHABILIDAD': classification.get('MICROHABILIDAD'),
                'COMPETENCIA MICROHABILIDAD': classification.get('COMPETENCIA MICROHABILIDAD'),
                'Contexto': contexto,
                'Enunciado': enunciado,
                'Opcion_A': opcion_a_match.group(1).strip() if opcion_a_match else "",
                'Opcion_B': opcion_b_match.group(1).strip() if opcion_b_match else "",
                'Opcion_C': opcion_c_match.group(1).strip() if opcion_c_match else "",
                'Opcion_D': opcion_d_match.group(1).strip() if opcion_d_match else "",
                'Clave': clave_match.group(1).strip() if clave_match else "",
                'Justificacion_A': just_a,
                'Justificacion_B': just_b,
                'Justificacion_C': just_c,
                'Justificacion_D': just_d,
            }
            datos_para_excel.append(fila)
    
        if not datos_para_excel:
            return None
    
        df = pd.DataFrame(datos_para_excel)
        
        column_order = ['item', 'ID', 'GRADO', 'ÁREA', 'ASIGNATURA', 'MACROHABILIDAD', 'PROCESO COGNITIVO', 
                        'MICROHABILIDAD', 'COMPETENCIA MICROHABILIDAD',
                        'Contexto', 'Enunciado', 'Opcion_A', 'Opcion_B', 'Opcion_C', 'Opcion_D', 'Clave',
                        'Justificacion_A', 'Justificacion_B', 'Justificacion_C', 'Justificacion_D']
        df = df[column_order]
    
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='Items')
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                writer.sheets['Items'].set_column(df.columns.get_loc(column), df.columns.get_loc(column), min(column_length, 50))
    
        buffer.seek(0)
        return buffer

    # --- LÓGICA PRINCIPAL DE LA APLICACIÓN ---
    
    # 1. Carga de datos desde Google Cloud Storage
    GCS_BUCKET_NAME = os.environ.get("GCS_BUCKET_NAME") 
    GCS_EXCEL_PATH = os.environ.get("GCS_EXCEL_PATH")
    GCS_PDF_PATH = os.environ.get("GCS_PDF_PATH")

    st.sidebar.header("Fuente de Datos (GCS)")
    
    df_datos = None
    manual_reglas_texto = ""

    if GCS_BUCKET_NAME and GCS_EXCEL_PATH:
        df_datos = leer_excel_desde_gcs(GCS_BUCKET_NAME, GCS_EXCEL_PATH)
    else:
        st.sidebar.error("Variables de entorno para GCS no configuradas.")
        st.info("La aplicación requiere GCS_BUCKET_NAME y GCS_EXCEL_PATH para cargar los datos.")
        st.stop()

    if GCS_BUCKET_NAME and GCS_PDF_PATH:
        manual_reglas_texto = leer_pdf_desde_gcs(GCS_BUCKET_NAME, GCS_PDF_PATH)
        max_manual_length = 15000
        if len(manual_reglas_texto) > max_manual_length:
            st.sidebar.warning(f"Manual truncado a {max_manual_length} caracteres.")
            manual_reglas_texto = manual_reglas_texto[:max_manual_length]
        st.sidebar.info(f"Manual de reglas cargado ({len(manual_reglas_texto)} caracteres).")

    # --- AÑADIR ESTE BLOQUE ---
    st.sidebar.markdown("---")
    st.sidebar.header("Recurso Opcional (Libro)")
    
    # Variable de sesión para el nombre del libro
    if 'processed_pdf_name' not in st.session_state:
        st.session_state.processed_pdf_name = None

    pdf_itinerario = st.sidebar.file_uploader("Subir PDF del Itinerario/Libro", type="pdf")
    #
    if pdf_itinerario:
            # Si se sube un nuevo libro O es un libro diferente al procesado
            if pdf_itinerario.name != st.session_state.processed_pdf_name:
                
                # 1. Primero, le decimos a Streamlit que TODO lo siguiente va en la barra lateral
                with st.sidebar:
                    # 2. AHORA llamamos al spinner normal, y aparecerá en la barra lateral
                    with st.spinner(f"Procesando '{pdf_itinerario.name}'... Esto puede tardar unos minutos."):
                
                        # --- INICIO DE LA SECCIÓN CORREGIDA ---
                        # Todo lo que sigue debe estar indentado a este nivel
                        # para que ocurra DENTRO del 'with st.spinner'
    
                        # 1. Borramos el índice viejo (si existe)
                        if 'pdf_index' in st.session_state:
                            del st.session_state['pdf_index']
                    
                        # 2. Procesamos el nuevo libro
                        pdf_bytes = pdf_itinerario.getvalue()
                        texto_completo = extraer_texto_pdf(pdf_bytes)
                        
                        if texto_completo:
                            # 3. Dividir (Chunking)
                            text_splitter = RecursiveCharacterTextSplitter(
                                chunk_size=1000, # 1000 caracteres por pedazo
                                chunk_overlap=100  # 100 caracteres de superposición
                            )
                            chunks = text_splitter.split_text(texto_completo)
                            
                            # 4. Vectorizar (Embedding) y almacenar
                            st.session_state['pdf_index'] = crear_indice_vectorial(chunks)
                            st.session_state.processed_pdf_name = pdf_itinerario.name
                            
                            # Mostramos el éxito (aún dentro del 'with st.sidebar')
                            st.sidebar.success(f"Libro '{pdf_itinerario.name}' procesado. {len(chunks)} secciones indexadas.")
                        else:
                            st.sidebar.error("El PDF está vacío o no se pudo leer.")
                        # --- FIN DE LA SECCIÓN CORREGIDA ---
    
            # Si el libro es el mismo que ya está cargado, no hacemos nada
            # y solo mostramos el mensaje de éxito.
            elif 'pdf_index' in st.session_state:
                 st.sidebar.success(f"Libro '{pdf_itinerario.name}' listo.")
                 
    # 2. Lógica de Generación y Auditoría de Ítems
    st.header("Generación y Auditoría de Ítems.")
    
    if df_datos is None:
        st.error("No se pudo cargar el archivo Excel desde GCS. Verifica la configuración.")
    else:
        # --- INTERFAZ DE USUARIO ---
        st.subheader("1. Selecciona los Criterios para la Generación")

        all_grades = df_datos['GRADO'].dropna().unique().tolist()
        grado_seleccionado = st.selectbox("Grado", sorted(all_grades), key="grado_sel")

        df_filtrado_grado = df_datos[df_datos['GRADO'].astype(str).str.upper() == str(grado_seleccionado).upper()]
        all_areas = df_filtrado_grado['ÁREA'].dropna().unique().tolist()
        area_seleccionada = st.selectbox("Área", sorted(all_areas), key="area_sel")

        df_filtrado_area = df_filtrado_grado[df_filtrado_grado['ÁREA'].astype(str).str.upper() == str(area_seleccionada).upper()]
        all_asignaturas = df_filtrado_area['ASIGNATURA'].dropna().unique().tolist()
        asignatura_seleccionada = st.selectbox("Asignatura", sorted(all_asignaturas), key="asignatura_sel")

        df_filtrado_asignatura = df_filtrado_area[df_filtrado_area['ASIGNATURA'].astype(str).str.upper() == str(asignatura_seleccionada).upper()]
        all_macrohabilidades = df_filtrado_asignatura['MACROHABILIDAD'].dropna().unique().tolist()
        macrohabilidad_seleccionada = st.selectbox("Macrohabilidad", sorted(all_macrohabilidades), key="macrohabilidad_sel")
    
        # --- LÓGICA DE CARGA Y GESTIÓN DE PROGRESO (CORREGIDA) ---
        nombre_archivo_progreso = generar_nombre_archivo_progreso(grado_seleccionado, asignatura_seleccionada, macrohabilidad_seleccionada)
        
        # Verificamos si la lista de ítems aprobados no existe en la sesión y la creamos vacía.
        # ESTA ES LA LÍNEA CLAVE QUE SOLUCIONA EL ERROR.
        if 'approved_items' not in st.session_state:
            st.session_state['approved_items'] = []
        
        # Ahora, gestionamos la carga de progreso si el usuario cambia de macrohabilidad.
        if st.session_state.get('current_macrohabilidad') != macrohabilidad_seleccionada:
            # Si la macrohabilidad cambió, cargamos el progreso desde GCS.
            # Esto reemplazará la lista vacía con los ítems guardados, si existen.
            st.session_state['approved_items'] = cargar_progreso_desde_gcs(GCS_BUCKET_NAME, nombre_archivo_progreso)
            st.session_state['current_review_index'] = 0
            
        # Finalmente, actualizamos la macrohabilidad y nombre de archivo actuales en la sesión.
        st.session_state['current_macrohabilidad'] = macrohabilidad_seleccionada
        st.session_state['nombre_archivo_progreso'] = nombre_archivo_progreso
        # --- FIN DE LA LÓGICA ---
        
        df_filtrado_macrohabilidad = df_filtrado_asignatura[df_filtrado_asignatura['MACROHABILIDAD'].astype(str).str.upper() == str(macrohabilidad_seleccionada).upper()]
        
        # --- ORDEN CORREGIDO: SECCIÓN MOVIDA HACIA ARRIBA ---
        st.markdown("---")
        st.subheader("2. Configuración de Modelos de Vertex AI")
        vertex_ai_models = [
            "gemini-2.5-pro",
            "gemini-2.5-flash",
            "gemini-2.5-flash-lite"
        ]
        col1, col2 = st.columns(2)
        with col1:
            gen_model_name = st.selectbox("**Modelo para Generación**", vertex_ai_models, index=1, key="gen_vertex_name")
        with col2:
            audit_model_name = st.selectbox("**Modelo para Auditoría**", vertex_ai_models, index=0, key="audit_vertex_name")
        # --- FIN DEL ORDEN CORREGIDO ---
        
        st.markdown("---")
        st.subheader("3. Selecciona las Habilidades y la Cantidad de Ítems")

        # --- INICIO DE LA NUEVA LÓGICA DE SELECCIÓN MÚLTIPLE ---

        # Preparamos la lista de habilidades para mostrar en la interfaz
        df_habilidades = df_filtrado_macrohabilidad.drop_duplicates(subset=['MICROHABILIDAD']).reset_index(drop=True)
        
        # Guardamos el dataframe de habilidades en el estado de la sesión para usarlo después
        st.session_state['df_habilidades_macrohabilidad'] = df_habilidades

        # Creamos un diccionario para guardar las selecciones del usuario: {indice: cantidad}
        if 'selecciones_usuario' not in st.session_state:
            st.session_state['selecciones_usuario'] = {}

        st.info("Marca las casillas de las habilidades que deseas generar y elige cuántos ítems necesitas para cada una.")
        
        # --- CÓDIGO DE REEMPLAZO ---
        
        # --- Lógica de Contexto General (Opcional y Corregida Definitivamente) ---
        contexto_general_macrohabilidad = ""
        with st.expander("📝 Opcional: Generar un contexto general para la macrohabilidad"):
            
            # 1. Inicializamos la variable de estado principal si no existe. Esta será nuestra ÚNICA fuente de verdad.
            if 'generated_context' not in st.session_state:
                st.session_state.generated_context = ""
            if 'show_context_refinement' not in st.session_state:
                st.session_state.show_context_refinement = False
        
            # --- WIDGETS DE SELECCIÓN (Sin cambios) ---
            categorias_contexto = ["No usar contexto general", "Contexto Escolar", "Contexto Cotidiano", "Contexto Científico", "Contexto Histórico", "Contexto Literario", "Contexto Político/Social", "Contexto Tecnológico", "Fragmento para Lectura", "Otro..."]
            categoria_elegida = st.selectbox("Elige un tipo de contexto:", categorias_contexto, key="ctx_categoria")
            tipo_contexto_final = categoria_elegida
            if categoria_elegida == "Fragmento para Lectura":
                tipos_fragmento = ["Crónica", "Noticia", "Entrevista", "Ensayo", "Cuento Corto", "Manual"]
                tipo_contexto_final = st.selectbox("Elige el tipo de fragmento:", tipos_fragmento, key="ctx_fragmento")
            elif categoria_elegida == "Otro...":
                tipo_contexto_final = st.text_input("Especifica el tipo de contexto que deseas:", key="ctx_otro", placeholder="Ej: Contexto mitológico griego")
            idea_usuario_ctx = st.text_area("Opcional: Da una idea o borrador para guiar a la IA...", key="ctx_idea", placeholder="Ej: Un equipo de biólogos marinos descubre una nueva especie...")
            
            if categoria_elegida != "No usar contexto general":
                if st.button("🧠 Generar Contexto con IA", key="btn_gen_ctx"):
                    with st.spinner("Generando contexto..."):
                        contexto_sugerido = generar_contexto_general_con_llm(gen_model_name, grado_seleccionado, area_seleccionada, asignatura_seleccionada, macrohabilidad_seleccionada, tipo_contexto=tipo_contexto_final, idea_usuario=idea_usuario_ctx)
                        if contexto_sugerido:
                            # Al generar, actualizamos directamente nuestra única variable de estado.
                            st.session_state.generated_context = contexto_sugerido
                            st.session_state.show_context_refinement = False
                            st.rerun()
        
            # --- EDICIÓN Y REFINAMIENTO (LÓGICA SIMPLIFICADA) ---
            if st.session_state.generated_context:
                st.markdown("---")
                st.markdown("##### Contexto Generado (puedes editarlo directamente):")
                
                # CAMBIO CLAVE #1: El widget ahora usa la misma clave que nuestra variable de estado.
                # Cualquier edición del usuario actualizará AUTOMÁTICAMENTE st.session_state.generated_context.
                st.text_area(
                    "Contexto generado",
                    key="generated_context",  # Esta es la clave unificada
                    height=200,
                    label_visibility="collapsed"
                )
                
                if st.button("✍️ Refinar Contexto con Feedback", key="btn_show_refine_ctx"):
                    st.session_state.show_context_refinement = not st.session_state.get('show_context_refinement', False)
                    st.rerun()
        
                if st.session_state.get('show_context_refinement', False):
                    with st.form("refine_context_form"):
                        feedback_ctx = st.text_area("Escribe tus observaciones para refinar:", key="ctx_feedback")
                        submitted = st.form_submit_button("🔄 Refinar con estas Observaciones")
                        
                        if submitted and feedback_ctx:
                            # CAMBIO CLAVE #2: Leemos directamente del estado principal, que gracias al
                            # cambio #1, ya contiene las ediciones manuales del usuario.
                            contexto_base_actual = st.session_state.generated_context
                            
                            with st.spinner("Refinando contexto con tu feedback..."):
                                contexto_refinado = refinar_contexto_con_llm(
                                    gen_model_name,
                                    contexto_original=contexto_base_actual,
                                    feedback_usuario=feedback_ctx
                                )
                                
                                if contexto_refinado:
                                    # CAMBIO CLAVE #3: Solo necesitamos actualizar nuestra única variable.
                                    # El widget se actualizará solo en el rerun.
                                    st.session_state.generated_context = contexto_refinado
                                    st.session_state.show_context_refinement = False
                                    st.rerun()
                                else:
                                    st.error("No se pudo refinar el contexto.")
        
            # La variable final simplemente lee el estado principal, que siempre está actualizado.
            contexto_general_macrohabilidad = st.session_state.get('generated_context', "").strip()


        descripcion_imagen_aprobada = ""
        with st.expander("🖼️ Opcional: Usar una imagen como insumo para el ítem"):
            
            # Inicializamos el estado si no existe
            if 'descripcion_imagen' not in st.session_state:
                st.session_state['descripcion_imagen'] = ""
        
            uploaded_file = st.file_uploader(
                "Sube un archivo de imagen (PNG, JPG) o un PDF de una sola página",
                type=['png', 'jpg', 'jpeg', 'pdf']
            )
            
            if uploaded_file:
                # Guardamos los bytes de la imagen en la sesión para usarla después
                st.session_state['source_image_bytes'] = uploaded_file.getvalue()
                st.session_state['source_image_type'] = uploaded_file.type
            else:
                # Si no hay archivo, nos aseguramos de que no haya una imagen vieja en memoria
                if 'source_image_bytes' in st.session_state:
                    del st.session_state['source_image_bytes']
                if 'source_image_type' in st.session_state:
                    del st.session_state['source_image_type']
                    
            if uploaded_file is not None:
                # Botón para activar el análisis
                if st.button("🧠 Analizar y Describir Imagen"):
                    with st.spinner("Analizando la imagen con IA..."):
                        file_bytes = uploaded_file.getvalue()
                        mime_type = uploaded_file.type
                        
                        # Usamos un modelo multimodal (Pro es ideal para esto)
                        descripcion_generada = describir_imagen_con_llm(
                            "gemini-2.5-flash", # O el modelo multimodal que prefieras
                            file_bytes, 
                            mime_type
                        )
                        
                        if descripcion_generada:
                            st.session_state['descripcion_imagen'] = descripcion_generada
                            st.rerun() # Refresca para mostrar el text_area
        
            # Si ya hay una descripción generada, la mostramos para edición
            if st.session_state['descripcion_imagen']:
                st.markdown("##### Descripción Generada (puedes editarla):")
                
                edited_description = st.text_area(
                    "Edita la descripción si es necesario:",
                    value=st.session_state['descripcion_imagen'],
                    height=250,
                    key="desc_img_edited"
                )
                
                # El texto final que se usará es el que esté en el área de texto
                descripcion_imagen_aprobada = edited_description.strip()
                st.success("✅ La descripción está lista para ser usada en la generación del ítem.")
        
        st.markdown("---") # Separador visual

        # Creamos la interfaz interactiva para la selección
        for index, row in df_habilidades.iterrows():
            proceso = row['PROCESO COGNITIVO']
            micro = row['MICROHABILIDAD']
            label = f"**{proceso}** // {micro}"
            
            # Usamos el índice como identificador único
            is_checked = st.checkbox(label, key=f"cb_{index}")

            if is_checked:
                # Si está marcado, mostramos el selector de cantidad y guardamos la elección
                cantidad = st.selectbox(
                    "Cantidad de ítems:",
                    options=[1, 2, 3],
                    key=f"qty_{index}",
                    label_visibility="collapsed"
                )
                st.session_state['selecciones_usuario'][index] = cantidad
            elif index in st.session_state['selecciones_usuario']:
                # Si se desmarca, lo eliminamos de las selecciones
                del st.session_state['selecciones_usuario'][index]

        # Definimos df_item_seleccionado como el dataframe completo para que la validación posterior funcione
        # La lógica real de selección se basa en 'selecciones_usuario'
        df_item_seleccionado = df_filtrado_macrohabilidad.copy()
        
        # --- FIN DE LA NUEVA LÓGICA DE SELECCIÓN MÚLTIPLE ---
        
        if df_item_seleccionado is None or df_item_seleccionado.empty:
            st.error("No hay datos válidos para generar ítems con los filtros actuales.")
        else:
            # --- El código de prompts y el botón ahora están dentro del 'else' ---
            st.markdown("---")
            st.subheader("4. Personaliza con Prompts Adicionales (Opcional)")
            prompt_bloom_adicional, prompt_construccion_adicional, prompt_especifico_adicional, prompt_auditor_adicional = "", "", "", ""
            if st.checkbox("Activar Prompts Adicionales"):
                st.info("Estos prompts se añadirán a las instrucciones de la IA para un control más fino.")
                prompt_bloom_adicional = st.text_area("Prompt para Taxonomía de Bloom:", help="Ej: 'Asegúrate que la pregunta requiera que el estudiante compare dos conceptos...'")
                prompt_construccion_adicional = st.text_area("Prompt para Construcción de Ítem:", help="Ej: 'Usa un lenguaje formal y evita coloquialismos.'")
                prompt_especifico_adicional = st.text_area("Prompt para Consideraciones Específicas:", help="Ej: 'El contexto debe estar relacionado con la ecología de un bosque.'")
                prompt_auditor_adicional = st.text_area("Prompt para el Auditor:", help="Ej: 'Verifica que la dificultad sea adecuada para un examen final.'")

            st.markdown("---")
            
            # =============================================================================
            # BLOQUE 1 MODIFICADO: LÓGICA DEL BOTÓN PRINCIPAL
            # =============================================================================
            if st.button("🚀 Generar y Auditar Ítem(s)"):
                if not st.session_state.get('selecciones_usuario'):
                    st.warning("⚠️ Por favor, selecciona al menos una habilidad para generar ítems.")
                else:
                    criterios_para_preguntas = {
                        "tipo_pregunta": "opción múltiple con 4 opciones",
                        "dificultad": "media",
                        "contexto_educativo": "estudiantes Colombianos entre 10 y 17 años",
                    }

                    # --- Construimos la nueva "cola de tareas" basada en la selección del usuario ---
                    items_para_procesar = []
                    df_habilidades_macrohabilidad = st.session_state['df_habilidades_macrohabilidad']

                    for index, cantidad in st.session_state.selecciones_usuario.items():
                        habilidad_seleccionada = df_habilidades_macrohabilidad.loc[index].to_dict()
                        for _ in range(cantidad):
                            items_para_procesar.append(habilidad_seleccionada)
                    
                    if items_para_procesar:
                        st.session_state.items_para_procesar = items_para_procesar
                        st.session_state.current_review_index = 0
                        st.session_state.awaiting_review = True
                        st.session_state.modo_lote = True
                        st.session_state.selecciones_usuario = {}
                        st.rerun()
            
            # =============================================================================
            # BLOQUE 2 MODIFICADO: SECCIÓN DE REVISIÓN CON GENERACIÓN SECUENCIAL Y GRÁFICOS
            # =============================================================================
            if 'awaiting_review' in st.session_state and st.session_state['awaiting_review']:
                
                current_index = st.session_state.get('current_review_index', 0)
                
                # --- LÓGICA DE CONTROL CON SESSION_STATE ---
                items_pendientes = st.session_state.get('items_para_procesar', [])
                total_items = len(items_pendientes)
                
                if current_index >= total_items:
                    st.session_state['awaiting_review'] = False
                    if 'item_under_review' in st.session_state:
                        del st.session_state['item_under_review']
                    st.rerun()
                else:
                    # Intentamos obtener el ítem de la memoria de la sesión
                    item_to_review = st.session_state.get('item_under_review')
                    
                    # Si NO hay un ítem en memoria, lo generamos
                    if item_to_review is None:
                        # 1. Crea un marcador de posición que ocupará un espacio en la pantalla
                        placeholder = st.empty()
            
                        # 2. Carga y muestra la animación DENTRO del marcador de posición
                        lottie_url = "https://lottie.host/41f1128a-22f4-40ad-99c8-1076328efb3e/MMre1fyJsg.json" # URL del dinosaurio
                        lottie_json = load_lottieurl(lottie_url)
            
                        with placeholder.container():
                            st.subheader(f"📝 Generando y Revisando Ítem ({current_index + 1} de {total_items})")
                            if lottie_json:
                                st_lottie(lottie_json, height=200, key="lottie_loading")
                            else:
                                st.info("Cargando animación...") # Mensaje de respaldo
            
                        # 3. EJECUTA TU PROCESO LARGO (La generación de la IA)
                        item_spec_row = items_pendientes[current_index]
                        current_fila_datos = {
                            'GRADO': grado_seleccionado, 'ÁREA': area_seleccionada, 'ASIGNATURA': asignatura_seleccionada, 'MACROHABILIDAD': macrohabilidad_seleccionada,
                            **item_spec_row
                        }
                        criterios_para_preguntas = {"tipo_pregunta": "opción múltiple con 4 opciones", "dificultad": "media", "contexto_educativo": "estudiantes Colombianos entre 10 y 17 años"}
                      
                        # --- AÑADIR ESTE BLOQUE DE BÚSQUEDA ---
                        contexto_del_libro = ""
                        if 'pdf_index' in st.session_state:
                            with st.spinner("Buscando en el libro guía..."):
                                # Usamos la microhabilidad como consulta
                                query_microhabilidad = current_fila_datos.get('MICROHABILIDAD', '')
                                
                                # Buscamos los 3 chunks más relevantes
                                chunks_relevantes = buscar_en_indice(query_microhabilidad, k=3)
                                
                                if chunks_relevantes:
                                    contexto_del_libro = "\n\n---\n\n".join(chunks_relevantes)
                                    st.info("ℹ️ Contexto relevante encontrado en el libro guía.")
                        # --- FIN DEL BLOQUE DE BÚSQUEDA ---

    
                        item_to_review = generar_pregunta_con_seleccion(
                            gen_model_name, audit_model_name, fila_datos=current_fila_datos,
                            criterios_generacion=criterios_para_preguntas, manual_reglas_texto=manual_reglas_texto,
                            contexto_general_macrohabilidad=contexto_general_macrohabilidad,
                            prompt_bloom_adicional=prompt_bloom_adicional, prompt_construccion_adicional=prompt_construccion_adicional,
                            prompt_especifico_adicional=prompt_especifico_adicional, prompt_auditor_adicional=prompt_auditor_adicional, descripcion_imagen_aprobada=descripcion_imagen_aprobada
                        )
                        st.session_state['item_under_review'] = item_to_review
                        
                        # 4. Una vez que el proceso termina, limpia el marcador de posición
                        placeholder.empty()
            
                    # Si se generó o encontró un ítem, lo mostramos para revisión
                    if item_to_review:
                        with st.expander("Ver detalles de clasificación del ítem", expanded=False):
                            st.json(item_to_review['classification'])
                        
                        st.markdown("##### Ítem Generado:")
                        st.text_area("Ítem", value=item_to_review['item_text'], height=400, key=f"item_text_{current_index}", disabled=True)
                        
                        st.markdown("##### Resultado de la Auditoría:")
                        status = item_to_review['final_audit_status']
                        if "✅" in status:
                            st.success(f"**Dictamen:** {status}")
                        elif "⚠️" in status:
                            st.warning(f"**Dictamen:** {status}")
                        else:
                            st.error(f"**Dictamen:** {status}")
                        st.markdown(f"**Observaciones:** {item_to_review['final_audit_observations']}")

                        # =============================================================================
                        # --- INICIO: INTEGRACIÓN DEL GENERADOR DE GRÁFICOS (MODIFICADO) ---
                        # =============================================================================

                        if item_to_review.get("grafico_necesario") == 'SÍ':
                            # Obtenemos la LISTA de descripciones que ya procesamos en el paso 2
                            descripciones = item_to_review.get("descripciones_graficos", [])
                            
                            # Usamos un expander de Streamlit para mostrar todos los gráficos de forma ordenada
                            with st.expander(f"🎨 Gráficos Requeridos para este Ítem ({len(descripciones)})", expanded=True):
                                
                                if not descripciones:
                                    st.warning("El ítem indica que requiere gráficos, pero no se encontró una descripción válida.")
                                else:
                                    # Iteramos sobre cada descripción de gráfico en la lista
                                    for idx, desc_grafico in enumerate(descripciones):
                                        
                                        # Extraemos la ubicación para etiquetar cada sección del gráfico
                                        ubicacion = desc_grafico.get("ubicacion", f"Gráfico #{idx + 1}").replace("_", " ").title()
                                        st.markdown(f"--- \n**Gráfico para:** `{ubicacion}`")
                        
                                        # Convertimos el diccionario del gráfico a un string JSON formateado para mostrarlo y editarlo.
                                        # 'ensure_ascii=False' es importante para que muestre bien las tildes y caracteres en español.
                                        descripcion_actual_str = json.dumps(desc_grafico, indent=2, ensure_ascii=False)
                                        
                                        # Clave ÚNICA para cada widget, ¡esto es muy importante para que Streamlit funcione bien en bucles!
                                        key_base = f"chart_{current_index}_{idx}"
                        
                                        edited_description = st.text_area(
                                            "Descripción JSON del Gráfico (puedes editarla):",
                                            value=descripcion_actual_str,
                                            key=f"desc_{key_base}", # Clave única para el área de texto
                                            height=200
                                        )
                        
                                        if st.button("🖼️ Generar / Actualizar Gráfico", key=f"btn_{key_base}"):
                                            if edited_description:
                                                
                                                # --- INICIO DE LA LÓGICA INTELIGENTE ---
                                                buffer_imagen = None
                                                
                                                # 1. Intentamos leer el texto como un JSON estructurado.
                                                try:
                                                    datos_grafico = json.loads(edited_description)
                                                    tipo_elemento = datos_grafico.get("tipo_elemento")
                                                except json.JSONDecodeError:
                                                    # Si falla (porque es texto simple), es un trabajo para el "Artista".
                                                    tipo_elemento = "otro_tipo"
                                                    datos_grafico = {} # Creamos un dict vacío para evitar errores

                                                # 2. Decidimos qué función llamar.
                                                if tipo_elemento == "otro_tipo":
                                                    # Es un trabajo creativo, llamamos al Artista (IA generativa).
                                                    with st.spinner("🤖 Creando visualización con IA generativa..."):
                                                        # Usamos la descripción natural del JSON o el texto completo si no es JSON.
                                                        prompt_para_imagen = datos_grafico.get("datos", {}).get("descripcion_natural", edited_description)
                                                        buffer_imagen = generar_imagen_con_ia(prompt_para_imagen)
                                                else:
                                                    # Es un trabajo de datos, llamamos al Ingeniero (el plugin de gráficos).
                                                    with st.spinner("⚙️ Construyendo gráfico desde datos..."):
                                                        _, buffer_imagen = generar_grafico_desde_texto(descripcion=edited_description)
                                                        
                                                # --- FIN DE LA LÓGICA INTELIGENTE ---

                                                # 3. Guardamos el resultado en la sesión.
                                                if buffer_imagen:
                                                    st.session_state[f'img_{key_base}'] = buffer_imagen
                                                    st.session_state[f'caption_{key_base}'] = f"Gráfico para '{ubicacion}' generado."
                                                else:
                                                    st.session_state[f'img_{key_base}'] = None
                                                    st.session_state[f'caption_{key_base}'] = "No se pudo generar el gráfico con la descripción proporcionada."

                                                # Forzamos un refresco para mostrar el resultado inmediatamente.
                                                st.rerun()
                                            else:
                                                st.warning("La descripción está vacía. No se puede generar el gráfico.")
                                        
                                        # Mostramos la imagen si ya fue generada y guardada en la sesión.
                                        if f'img_{key_base}' in st.session_state and st.session_state[f'img_{key_base}']:
                                            st.image(
                                                st.session_state[f'img_{key_base}'],
                                                caption=st.session_state.get(f'caption_{key_base}'),
                                                use_column_width=True
                                            )
                                        # Si no hay imagen pero sí un mensaje (porque falló), lo mostramos.
                                        elif f'caption_{key_base}' in st.session_state:
                                            st.warning(st.session_state[f'caption_{key_base}'])
                                            
                        # =============================================================================
                        # --- FIN: NUEVA INTERFAZ PARA MÚLTIPLES GRÁFICOS ---
                        # =============================================================================
                        # =============================================================================
                        # BLOQUE 3: LÓGICA DE BOTONES (VERIFICADA, CON CAMBIOS NECESARIOS)
                        # =============================================================================
                        col_aprob, col_rechazo, col_descartar = st.columns(3)
            
                        with col_aprob:
                            if st.button("👍 Aprobar y Siguiente", key=f"approve_{current_index}", use_container_width=True):
                                
                                # --- INICIO: NUEVA LÓGICA PARA CAPTURAR IMÁGENES ---
                                if 'source_image_bytes' in st.session_state:
                                        item_to_review['source_image'] = io.BytesIO(st.session_state['source_image_bytes'])
                                if item_to_review.get("grafico_necesario") == 'SÍ':
                                    generated_images = []
                                    descripciones = item_to_review.get("descripciones_graficos", [])
                                    for idx, desc in enumerate(descripciones):
                                        key_base = f"chart_{current_index}_{idx}"
                                        # Buscamos si la imagen existe en la sesión
                                        if f'img_{key_base}' in st.session_state and st.session_state[f'img_{key_base}']:
                                            generated_images.append({
                                                "ubicacion": desc.get("ubicacion"),
                                                "buffer": st.session_state[f'img_{key_base}']
                                            })
                                    # Añadimos la lista de imágenes capturadas al ítem
                                    item_to_review['generated_images'] = generated_images
                                # --- FIN: NUEVA LÓGICA ---
                        
                                st.session_state.approved_items.append(item_to_review)
                                guardar_progreso_en_gcs(GCS_BUCKET_NAME, st.session_state.nombre_archivo_progreso, st.session_state.approved_items)
                                
                                # ... (el resto del código para limpiar la sesión y hacer rerun se mantiene igual) ...
                                # Limpiamos las imágenes de la sesión para el siguiente ítem
                                for idx, desc in enumerate(item_to_review.get("descripciones_graficos", [])):
                                    key_base = f"chart_{current_index}_{idx}"
                                    if f'img_{key_base}' in st.session_state: del st.session_state[f'img_{key_base}']
                                    if f'caption_{key_base}' in st.session_state: del st.session_state[f'caption_{key_base}']
                        
                                st.session_state.current_review_index += 1
                                st.session_state['show_feedback_form'] = False
                                st.session_state['item_under_review'] = None
                                st.rerun()
                                
                        with col_rechazo:
                            if st.button("✍️ Refinar con Feedback", key=f"refine_{current_index}", use_container_width=True):
                                # Hacemos el cambio de estado de forma más explícita
                                if 'show_feedback_form' not in st.session_state:
                                    st.session_state.show_feedback_form = True
                                else:
                                    st.session_state.show_feedback_form = not st.session_state.show_feedback_form
                                
                                # Forzamos un rerun para asegurar que la app se actualice con el nuevo estado
                                st.rerun()            
            
                        with col_descartar:
                            if st.button("👎 Descartar Ítem", key=f"discard_{current_index}", use_container_width=True):
                                if f'generated_chart_image_{current_index}' in st.session_state:
                                    del st.session_state[f'generated_chart_image_{current_index}']
                                if f'generated_chart_caption_{current_index}' in st.session_state:
                                    del st.session_state[f'generated_chart_caption_{current_index}']
                                if f'chart_description_{current_index}' in st.session_state:
                                    del st.session_state[f'chart_description_{current_index}']                   
                                st.session_state.current_review_index += 1
                                st.session_state['show_feedback_form'] = False
                                st.session_state['item_under_review'] = None # Limpiamos para generar el siguiente
                                st.rerun()
                        
                        # Formulario de Feedback
                        if st.session_state.get('show_feedback_form', False):
                            with st.form(key='feedback_form'):
                                st.markdown("---")
                                st.markdown("#### Proporciona tus observaciones para refinar el ítem:")
                                feedback_usuario = st.text_area(
                                    "Escribe aquí tus correcciones o sugerencias...",
                                    key="feedback_text"
                                )
                                submitted = st.form_submit_button("🔄 Refinar con estas Observaciones")
                        
                                if submitted and feedback_usuario:
                                    # 1. Recuperamos el ítem actual desde el session_state AHORA MISMO.
                                    item_actual_para_refinar = st.session_state.get('item_under_review')
                        
                                    # 2. Verificamos que el ítem realmente existe antes de continuar.
                                    if item_actual_para_refinar and 'item_text' in item_actual_para_refinar:
                                        with st.spinner("🧠 Refinando el ítem con tu feedback..."):
                                            classif_norm = normaliza_claves_classif(item_actual_para_refinar.get('classification', {}))
                                            
                                            # 3. Llamamos a la función con los datos correctos y verificados.
                                            refined_item_data = generar_pregunta_con_seleccion(
                                                gen_model_name=st.session_state.gen_vertex_name,
                                                audit_model_name=st.session_state.audit_vertex_name,
                                                fila_datos=classif_norm,
                                                criterios_generacion={
                                                    "tipo_pregunta": "opción múltiple con 4 opciones", "dificultad": "media",
                                                    "contexto_educativo": "estudiantes Colombianos entre 10 y 17 años"
                                                },
                                                manual_reglas_texto=manual_reglas_texto,
                                                feedback_usuario=feedback_usuario,
                                                # Usamos el texto del ítem que acabamos de recuperar.
                                                item_a_refinar_text=item_actual_para_refinar['item_text']
                                            )
                        
                                            if refined_item_data:
                                                st.session_state['item_under_review'] = refined_item_data
                                                st.session_state['show_feedback_form'] = False
                                                st.success("¡Ítem refinado! Por favor, revísalo de nuevo.")
                                                st.rerun()
                                            else:
                                                st.error("Fallo al refinar el ítem. Intenta de nuevo o ajusta tu feedback.")
                                    else:
                                        # Si por alguna razón el ítem se perdió, informamos al usuario.
                                        st.error("Error de estado: No se encontró el ítem a refinar. Por favor, descarte este ítem y genere uno nuevo.")

                                    
                                    if refined_item_data:
                                        # Reemplazamos el ítem en revisión con la versión refinada
                                        st.session_state['item_under_review'] = refined_item_data
                                        st.session_state['show_feedback_form'] = False
                                        st.success("¡Ítem refinado! Por favor, revísalo de nuevo.")
                                        st.rerun()
                                    else:
                                        st.error("Fallo al refinar el ítem. Intenta de nuevo o ajusta tu feedback.")

            if 'approved_items' in st.session_state and st.session_state['approved_items']:
                if not st.session_state.get('awaiting_review', False):
                    st.markdown("---")
                    st.subheader(f"✅ Ítems Aprobados: {len(st.session_state.approved_items)}")
                    st.success("Todos los ítems seleccionados han sido procesados. Ahora puedes exportarlos.")

                    nombre_archivo_zip = f"items_{macrohabilidad_seleccionada.replace(' ', '_')}.zip"
                    
                    zip_buffer = exportar_a_zip(st.session_state.approved_items)
                    st.download_button(
                        label="📥 Descargar todos los Ítems Aprobados (.zip)",
                        data=zip_buffer,
                        file_name=nombre_archivo_zip,
                        mime="application/zip",
                        use_container_width=True
                    )
                                        
                    st.write("")
                    nombre_base = macrohabilidad_seleccionada.replace(' ', '_').lower()
                    excel_buffer = exportar_a_excel(st.session_state.approved_items, nombre_base)
                    nombre_archivo_excel = f"items_aprobados_{nombre_base}.xlsx"
                    if excel_buffer:
                        st.download_button(
                            label="📥 Descargar Ítems Aprobados (.xlsx)",
                            data=excel_buffer,
                            file_name=nombre_archivo_excel,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    
                    with st.expander("Ver y Descargar Prompts Utilizados"):
                        st.info("Aquí puedes descargar un archivo de texto con los prompts completos que se enviaron a la IA para generar y auditar los ítems aprobados.")
                        
                        combined_prompts_content = ""
                        for i, item_data in enumerate(st.session_state['approved_items']):
                            combined_prompts_content += f"--- ÍTEM APROBADO #{i+1} ---\n"
                            combined_prompts_content += f"Clasificación: {item_data.get('classification', {})}\n"
                            combined_prompts_content += "="*40 + "\n\n"
                            combined_prompts_content += f"--- PROMPT DE GENERACIÓN ---\n"
                            combined_prompts_content += f"{item_data.get('generation_prompt_used', 'No disponible')}\n\n"
                            combined_prompts_content += f"--- PROMPT DE AUDITORÍA ---\n"
                            combined_prompts_content += f"{item_data.get('auditor_prompt_used', 'No disponible')}\n\n"
                            combined_prompts_content += "#"*80 + "\n\n"
                        
                        st.download_button(
                            label="📥 Descargar Prompts (.txt)",
                            data=combined_prompts_content.encode('utf-8'),
                            file_name=f"prompts_{macrohabilidad_seleccionada.replace(' ', '_')}.txt",
                            mime="text/plain",
                            use_container_width=True
                        )
                    
                    st.markdown("---")
                    if st.button("✨ Reset: Borrar información y generar nuevo ítem", use_container_width=True, type="primary"):
                        
                        # --- LÍNEAS AÑADIDAS PARA BORRAR PROGRESO ---
                        if 'nombre_archivo_progreso' in st.session_state:
                            borrar_progreso_en_gcs(GCS_BUCKET_NAME, st.session_state.nombre_archivo_progreso)
                        # ----------------------------------------------
                    
                        # Limpiar todos los estados relevantes
                        
                        keys_to_pop = ['approved_items', 'processed_items_list_for_review', 'current_review_index',
                                       'awaiting_review', 'items_para_procesar', 'modo_lote', 'show_feedback_form',
                                       'context_approved', 'generated_context', 'show_context_options', 'nombre_archivo_progreso',
                                       'source_image_bytes', 'source_image_type', 'descripcion_imagen']

                        for key in keys_to_pop:
                            if key in st.session_state:
                                st.session_state.pop(key)
                        
                        st.rerun()

# --- BLOQUE DE EJECUCIÓN ---
if __name__ == "__main__":
    load_dotenv() # Carga las variables de entorno

    # --- AUTENTICACIÓN SEGURA CON VARIABLES DE ENTORNO ---
    # Lee la contraseña desde los secretos configurados en la plataforma de despliegue
    PASSWORD = os.environ.get("PASSWORD")

    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False

    if not st.session_state["authenticated"]:
        st.set_page_config(layout="centered")
        st.title("🔒 Acceso restringido")
        pwd = st.text_input("Introduce la contraseña:", type="password")

        if st.button("Entrar"):
            if pwd and pwd == PASSWORD:
                st.session_state["authenticated"] = True
                st.rerun()
            else:
                st.error("❌ Contraseña incorrecta")
        st.stop()
    

    # --- SI LA CONTRASEÑA ES CORRECTA, EJECUTA LA APP ---
    try:
        main()
    except Exception as e:
        st.set_page_config(layout="centered")
        st.title("🛑 Error Crítico de la Aplicación")
        st.error(
            "Ocurrió un error grave que impidió que la aplicación se iniciara correctamente. "
            "Esto suele ser un problema de configuración o permisos en Google Cloud."
        )
        st.subheader("Mensaje de Error Técnico:")
        st.exception(e)
